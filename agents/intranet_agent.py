from google.oauth2 import service_account
from googleapiclient.discovery import build
from services.Common_Functions import prune_state
from services.Graph_state import GraphState
from services.llm_service import call_llm
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import io
import zipfile
from html import unescape

import tempfile
from docx import Document
import os
import re
import streamlit as st
from dotenv import load_dotenv

#Sammple Query
# What is the underwriting framework for Aerospace liability policies

STATE_KEYS_SET_AT_ENTRY = []
load_dotenv()


# ==========================================
# LOB CONFIGURATION
# ==========================================

# ==========================================
# LOB CONFIGURATION
# ==========================================

LOB_KEYWORDS = {
    "Marine": [
        "marine", "cargo", "vessel", "ship", "shipping",
        "maritime", "freight", "tanker", "bulk carrier",
        "p&i", "protection and indemnity", "hull"
    ],

    "Casualty": [
        "liability", "general liability", "public liability",
        "employers liability", "third party liability",
        "bodily injury", "property damage", "casualty",
        "umbrella", "excess liability"
    ],

    "Property": [
        "property", "fire", "industrial all risk",
        "building", "contents", "asset", "stock",
        "warehouse", "plant", "machinery breakdown",
        "business interruption"
    ],

    "Motor": [
        "motor", "auto", "vehicle", "fleet",
        "car", "truck", "commercial vehicle",
        "third party motor", "own damage"
    ],

    "Energy": [
        "energy", "oil", "gas", "power",
        "renewable", "offshore", "onshore",
        "rig", "platform", "pipeline",
        "solar", "wind"
    ],

    "Aviation": [
        "aircraft", "aviation", "aero", "airplane",
        "flight", "pilot", "airline", "airport",
        "aerospace", "hangarkeepers", "uav",
        "drone", "charter operator","aerospace"
    ],

    "Financial Lines": [
        "directors and officers", "d&o",
        "professional indemnity", "e&o",
        "financial lines", "crime",
        "fidelity", "cyber", "management liability",
        "errors and omissions"
    ],

    "Construction": [
        "construction", "contractor",
        "builders risk", "construction all risk",
        "project", "infrastructure",
        "bridge", "tunnel", "civil works",
        "erection all risk"
    ]
}


LOB_FOLDERS = {
    "Marine": os.getenv("GOOGLE_DRIVE_MARINE_FOLDER_ID"),
    "Casualty": os.getenv("GOOGLE_DRIVE_CASUALTY_FOLDER_ID"),
    "Property": os.getenv("GOOGLE_DRIVE_PROPERTY_FOLDER_ID"),
    "Motor": os.getenv("GOOGLE_DRIVE_MOTOR_FOLDER_ID"),
    "Energy": os.getenv("GOOGLE_DRIVE_ENERGY_FOLDER_ID"),
    "Aviation": os.getenv("GOOGLE_DRIVE_AVIATION_FOLDER_ID") or os.getenv("GOOGLE_DRIVE_AERO_FOLDER_ID"),
    "Financial Lines": os.getenv("GOOGLE_DRIVE_FINANCIAL_LINES_FOLDER_ID"),
    "Construction": os.getenv("GOOGLE_DRIVE_CONSTRUCTION_FOLDER_ID"),
}


def _normalize_env_path(value: str | None) -> str | None:
    if not value:
        return None
    # Handle accidental escape-parsing artifacts from quoted .env values (e.g. "\a" -> bell char).
    cleaned = value.strip().strip('"').strip("'")
    # A bell char in a Windows path usually means "\a" was parsed; restore as "\a".
    cleaned = cleaned.replace("\x07", "\\a")
    return os.path.expandvars(os.path.expanduser(cleaned))


def _load_service_account_info_from_secrets() -> dict | None:
    """
    Load Google service-account credentials from Streamlit secrets.
    Supported secrets:
    1) GOOGLE_DRIVE_CREDENTIALS_JSON = "{...json...}"
    2) [google_service_account] table with credential fields
    """
    try:
        # Flat JSON string secret
        raw = st.secrets.get("GOOGLE_DRIVE_CREDENTIALS_JSON")
        if raw:
            if isinstance(raw, str):
                return json.loads(raw)
            return dict(raw)

        # Structured table secret
        table = st.secrets.get("google_service_account")
        if table:
            return dict(table)
    except Exception:
        return None
    return None


def detect_lob_from_query(query: str) -> str:
    query_lower = query.lower()
    scores = {}

    for lob, keywords in LOB_KEYWORDS.items():
        score = 0
        for keyword in keywords:
            if keyword.lower() in query_lower:
                score += 1
        scores[lob] = score

    best_lob = max(scores, key=scores.get)
    best_score = scores[best_lob]

    if best_score == 0:
        st.error("Please select a Line of Business (Aero, Marine, Construction).")
        st.stop()

    # st.success(f"Detected {best_lob}")
    return best_lob



# ==========================================
# INTRANET AGENT
# ==========================================

class IntranetAgent:

    def __init__(self, credentials_path: str = None):
        self.credentials_path = credentials_path or _normalize_env_path(os.getenv("GOOGLE_DRIVE_CREDENTIALS_PATH"))
        self.service = self._authenticate()
        self.embeddings = OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY"))
        self.policy_index = None

    def _authenticate(self):
        try:
            scopes = ['https://www.googleapis.com/auth/drive.readonly']
            secret_info = _load_service_account_info_from_secrets()

            if secret_info:
                credentials = service_account.Credentials.from_service_account_info(
                    secret_info,
                    scopes=scopes,
                )
                return build('drive', 'v3', credentials=credentials)

            if self.credentials_path and os.path.exists(self.credentials_path):
                credentials = service_account.Credentials.from_service_account_file(
                    self.credentials_path,
                    scopes=scopes,
                )
                return build('drive', 'v3', credentials=credentials)

            st.warning("Google Drive not configured. Add GOOGLE_DRIVE_CREDENTIALS_JSON in Streamlit secrets.")
            return None

        except json.JSONDecodeError:
            st.error("Invalid GOOGLE_DRIVE_CREDENTIALS_JSON in Streamlit secrets.")
            return None
        except Exception as e:
            st.error(f"Drive authentication failed: {e}")
            return None

    # --------------------------------------
    # SEARCH DOCUMENTS (LOB ROUTING)
    # --------------------------------------

    def search_documents(self, query: str, lob: str, max_results: int = 3):
        if not self.service:
            return []

        folder_id = LOB_FOLDERS.get(lob)
        if not folder_id:
            st.error(f"Folder ID missing for {lob}")
            return []

        safe_query = query.replace("'", "")

        search_query = (
            f"'{folder_id}' in parents and "
            f"(name contains '{safe_query}' or fullText contains '{safe_query}') "
            f"and trashed=false"
        )

        try:
            results = self.service.files().list(
                q=search_query,
                pageSize=max_results,
                fields="files(id, name, mimeType, createdTime, webViewLink)"
            ).execute()

            return results.get("files", [])

        except Exception as e:
            st.error(f"Drive search failed: {e}")
            return []
        
    # --------------------------------------
    # LIST ALL FILES IN LOB FOLDER
    # --------------------------------------
    def list_all_files_in_folder(self, lob: str):

        if not self.service:
            return []

        folder_id = LOB_FOLDERS.get(lob)

        if not folder_id:
            st.error(f"Folder ID missing for {lob}")
            return []

        try:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                pageSize=1000,
                fields="files(id, name, mimeType, webViewLink)"
            ).execute()

            return results.get("files", [])

        except Exception as e:
            st.error(f"Folder listing failed: {e}")
            return []


    # --------------------------------------
    # DOWNLOAD
    # --------------------------------------

    def download_document(self, file_id: str, mime_type: str):
        if not self.service:
            return None

        try:
            from googleapiclient.http import MediaIoBaseDownload
            import io

            file_content = io.BytesIO()

            #  If Google Docs / Sheets / Slides → export
            if mime_type.startswith("application/vnd.google-apps"):
                request = self.service.files().export(
                    fileId=file_id,
                    mimeType="text/plain"  # Best for RAG
                )
            else:
                # 🔹 Normal binary file (PDF, DOCX, TXT)
                request = self.service.files().get_media(fileId=file_id)

            downloader = MediaIoBaseDownload(file_content, request)

            done = False
            while not done:
                _, done = downloader.next_chunk()

            file_content.seek(0)
            return file_content.read()

        except Exception as e:
            st.error(f"Download failed: {e}")
            return None

    # --------------------------------------
    # EXTRACT DOCX
    # --------------------------------------
    def _extract_text_from_docx_xml(self, docx_path: str) -> str:
        """
        Fallback for DOCX files where key content is stored in text boxes/runs
        not surfaced by python-docx paragraph traversal.
        """
        xml_targets = {
            "word/document.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
        }
        lines = []

        try:
            with zipfile.ZipFile(docx_path) as archive:
                for name in archive.namelist():
                    if name in xml_targets or name.startswith("word/header") or name.startswith("word/footer"):
                        xml_text = archive.read(name).decode("utf-8", errors="ignore")
                        for raw in re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_text):
                            line = unescape(raw).strip()
                            if line:
                                lines.append(line)
        except Exception:
            return ""

        return "\n".join(lines)

    def extract_text_from_docx(self, content: bytes) -> str:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            doc = Document(tmp_path)
            lines = []

            # Body paragraphs
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    lines.append(p.text.strip())

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            lines.append(cell.text.strip())

            # Headers and footers
            for section in doc.sections:
                for p in section.header.paragraphs:
                    if p.text and p.text.strip():
                        lines.append(p.text.strip())
                for p in section.footer.paragraphs:
                    if p.text and p.text.strip():
                        lines.append(p.text.strip())

            # XML fallback to catch text boxes/shapes not exposed by python-docx
            xml_fallback = self._extract_text_from_docx_xml(tmp_path)
            if xml_fallback:
                lines.extend([ln.strip() for ln in xml_fallback.splitlines() if ln.strip()])

            # De-duplicate while preserving order
            deduped = []
            seen = set()
            for line in lines:
                key = re.sub(r"\s+", " ", line).strip().lower()
                if key and key not in seen:
                    seen.add(key)
                    deduped.append(line)

            os.unlink(tmp_path)
            return "\n".join(deduped)

        except Exception as e:
            st.error(f"DOCX extraction failed: {e}")
            return ""
        
        # --------------------------------------
    # EXTRACT PDF
    # --------------------------------------
    def extract_text_from_pdf(self, content: bytes) -> str:
        try:
            pdf_stream = io.BytesIO(content)
            reader = PdfReader(pdf_stream)

            text = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

            return "\n".join(text)

        except Exception as e:
            st.error(f"PDF extraction failed: {e}")
            return ""

    # --------------------------------------
    # BUILD VECTOR STORE
    # --------------------------------------

    def build_policy_vector_store(self, documents):

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=200
        )

        texts = []
        metadatas = []

        for doc in documents:
            chunks = splitter.split_text(doc["content"])
            for chunk in chunks:
                texts.append(chunk)
                metadatas.append({
                    "file_name": doc["name"],
                    "source": doc.get("webViewLink", "")
                })

        if texts:
            return FAISS.from_texts(texts, self.embeddings, metadatas=metadatas)

        return None

    # --------------------------------------
    # SEMANTIC SEARCH
    # --------------------------------------

    def semantic_search(self, query: str, k: int = 5):
        if not self.policy_index:
            return []

        results = self.policy_index.similarity_search(query, k=k)
        return [(doc.page_content, doc.metadata) for doc in results]


# ==========================================
# NODE
# ==========================================

def intranet_node(state: GraphState) -> GraphState:

    user_prompt = state.get("user_prompt", "")

    try:
        agent = IntranetAgent()
    except Exception as e:
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "intranet_summary": f"Initialization failed: {str(e)}"
        }

    #  Detect LOB
    lob = detect_lob_from_query(user_prompt)

    #  GET ALL FILES FROM FOLDER (NO SEARCH)
    folder_files = agent.list_all_files_in_folder(lob)

    if not folder_files:
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "intranet_summary": f"No documents found inside {lob} folder.",
            "intranet_lob": lob
        }

    documents_with_content = []

    #  DOWNLOAD + EXTRACT ALL FILES
    for doc in folder_files:

        content = agent.download_document(doc["id"], doc.get("mimeType", ""))
        if not content:
            continue

        mime_type = doc.get("mimeType", "")

        if "wordprocessingml" in mime_type:
            text_content = agent.extract_text_from_docx(content)

        elif mime_type.startswith("application/vnd.google-apps"):
            text_content = content.decode("utf-8", errors="ignore")

        elif "pdf" in mime_type:
            text_content = agent.extract_text_from_pdf(content)

        else:
            text_content = content.decode("utf-8", errors="ignore")

        if text_content.strip():
            documents_with_content.append({
                "name": doc["name"],
                "content": text_content,
                "webViewLink": doc.get("webViewLink", "")
            })

    if not documents_with_content:
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "intranet_summary": "Content extraction failed.",
            "intranet_lob": lob
        }

    #  Build Vector Store (Per Session Per LOB)
    # Versioned cache key to avoid stale vector indexes across extraction logic updates.
    cache_key = f"policy_index_v2_{lob}"

    if cache_key not in st.session_state:
        st.session_state[cache_key] = agent.build_policy_vector_store(documents_with_content)

    agent.policy_index = st.session_state[cache_key]

    #  Semantic Search
    results = agent.semantic_search(user_prompt, k=8)

    # Add targeted retrieval for write/not-write sections so key exclusions are not missed.
    targeted_queries = [
        f"{user_prompt} business we underwrite business we write covered accepted",
        f"{user_prompt} business we do not underwrite business we do not write excluded declined restricted not covered",
    ]

    for tq in targeted_queries:
        try:
            results.extend(agent.semantic_search(tq, k=5))
        except Exception:
            pass

    # De-duplicate semantic results while preserving order.
    dedup_results = []
    seen_result_keys = set()
    for chunk, meta in results:
        file_name = meta.get("file_name", "")
        key = (file_name, re.sub(r"\s+", " ", chunk).strip()[:500])
        if key in seen_result_keys:
            continue
        seen_result_keys.add(key)
        dedup_results.append((chunk, meta))
    results = dedup_results

    if not results:
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "intranet_summary": "No relevant content found.",
            "intranet_lob": lob
        }

    # Focused evidence windows from full extracted text to explicitly surface
    # "business we write / do not write" sections when present.
    def _collect_evidence_windows(docs, keywords, window_after: int = 8, max_windows: int = 8):
        snippets = []
        lowered_keywords = [k.lower() for k in keywords]

        for doc_item in docs:
            lines = [ln.strip() for ln in doc_item.get("content", "").splitlines() if ln.strip()]
            for i, line in enumerate(lines):
                line_low = line.lower()
                if any(k in line_low for k in lowered_keywords):
                    start = max(0, i - 1)
                    end = min(len(lines), i + window_after + 1)
                    snippet = "\n".join(lines[start:end])
                    snippets.append(f"[{doc_item.get('name', 'Document')}]\n{snippet}")
                    if len(snippets) >= max_windows:
                        return snippets
        return snippets

    write_snippets = _collect_evidence_windows(
        documents_with_content,
        ["business we underwrite", "business we write", "will underwrite", "covered", "accepted"]
    )
    not_write_snippets = _collect_evidence_windows(
        documents_with_content,
        [
            "business we do not underwrite",
            "business we do not write",
            "will not underwrite",
            "not covered",
            "excluded",
            "declined",
            "restricted",
        ]
    )

    semantic_context = "\n\n---\n\n".join([r[0] for r in results[:10]])

    focused_evidence_parts = []
    if write_snippets:
        focused_evidence_parts.append("### Focused Evidence: Business Written\n" + "\n\n".join(write_snippets[:4]))
    if not_write_snippets:
        focused_evidence_parts.append("### Focused Evidence: Business Not Written\n" + "\n\n".join(not_write_snippets[:4]))

    focused_evidence = "\n\n---\n\n".join(focused_evidence_parts)
    context = semantic_context if not focused_evidence else f"{semantic_context}\n\n---\n\n{focused_evidence}"

    answer_prompt = f"""
You are an expert underwriting assistant.

Use ONLY the context below. Do not use outside knowledge.
If details are not found, explicitly state "Not specified in the provided documents."

Context:
{context}

Question:
{user_prompt}

Return your answer in this exact structure using markdown:

## Document Summary
- Write one concise paragraph (around 90-140 words) summarizing the document guidance relevant to the question.

## Key Points
- Provide exactly 3 to 4 bullet points.
- Each bullet must be specific and action-oriented (underwriting/claims/policy wording implications).

## Business Written
- List the business lines, risks, or activities that the document explicitly says are written/covered/accepted.
- If none are explicitly stated, write: "Not specified in the provided documents."

## Business Not Written
- List the business lines, risks, or activities that the document explicitly says are not written/excluded/declined/restricted.
- If none are explicitly stated, write: "Not specified in the provided documents."

Rules:
- Be specific, avoid generic statements.
- Do not mention information that is not present in the context.
- Prefer short, clear bullets.
"""

    answer = call_llm(answer_prompt)

    source_details = list({
        (meta["file_name"], meta.get("source", ""))
        for _, meta in results
    })

    web_links = [s[1] for s in source_details if s[1]]

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "intranet_summary": answer,
        "intranet_sources": source_details,
        "intranet_doc_links": web_links,
        "intranet_doc_count": len(documents_with_content),
        "intranet_lob": lob
    }
