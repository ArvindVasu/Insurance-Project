# Summarize the attached broker submission document and fetch the internal loss history, web insights for international casualty lines

from __future__ import annotations

import io
import json
import os
import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:
        PdfReader = None  # type: ignore

from agents.document_agent import document_node
from agents.intranet_agent import intranet_node
from agents.vanna_agent import vanna_node
from config.global_variables import DB_PATH
from services.Common_Functions import get_schema_description
from services.Graph_state import GraphState
from services.llm_service import call_llm
from services.vanna_service import vanna_configure

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_EOI_TEMPLATE_PATH = PROJECT_ROOT / "Doc" / "Insurance_EOI_Form.docx"
PDF_EOI_TEMPLATE_PATH = PROJECT_ROOT / "Doc" / "Insurance_EOI_Form.pdf"
DEFAULT_EOI_DB_PATH = PROJECT_ROOT / "Underwriter_Data.db"
CHECKED_BOX = "\u2611"
UNCHECKED_BOX = "\u2610"

SCORE_BANDS = {
    "write": 50,
    "conditional": 70,
    "refer": 85,
}

METRIC_WEIGHTS = {
    "loss_quality_composite": 0.25,
    "loss_pattern_risk": 0.15,
    "revenue_scale_risk": 0.10,
    "geographic_spread_risk": 0.15,
    "risk_management_quality": 0.10,
    "external_risk": 0.10,
    "coverage_complexity": 0.10,
    "guideline_fit": 0.05,
}

LOSS_PATTERN_WEIGHTS = {
    "claims_frequency_percentile": 0.40,
    "severity_percentile": 0.40,
    "incurred_percentile": 0.20,
}

SQL_HIDDEN_OUTPUT_COLUMNS = {
    "total_incurred_loss",
    "loss_ratio_percentile_basis",
    "loss_ratio_source",
    "client_loss_ratio_source",
    "client_claims_frequency_source",
    "client_severity_source",
    "client_incurred_source",
    "lob_hint",
    "source_db",
}

WEB_RISK_BANDS = {
    "low_max": 30.0,
    "moderate_max": 60.0,
}

FIELD_LABELS = {
    "registered_address": "Registered Address",
    "city": "City",
    "state": "State",
    "country": "Country",
    "postal_code": "Postal Code",
    "phone_number": "Phone Number",
    "email_address": "Email Address",
    "website": "Website (if applicable)",
    "expected_sum_insured": "Expected Sum Insured / Coverage Amount",
    "pan_tax_id": "PAN / Tax ID / Business ID",
    "gst_number": "GST Number (if applicable)",
    "regulatory_licenses": "Regulatory Licenses / Certifications",
    "claims_history": "Claims History (last 3–5 years)",
}

FIELD_MAX_LEN = {
    "registered_address": 240,
    "city": 80,
    "state": 80,
    "country": 80,
    "postal_code": 24,
    "phone_number": 40,
    "email_address": 120,
    "website": 180,
    "expected_sum_insured": 80,
    "pan_tax_id": 80,
    "gst_number": 40,
    "regulatory_licenses": 120,
    "claims_history": 260,
}

NOISE_TOKENS = [
    "submission date",
    "client overview",
    "program objective",
    "loss history",
    "coverage requested",
    "value proposition",
    "required deliverables",
    "disclaimer",
    "this submission",
]

DECLARATION_TEXT = (
    "I/We hereby declare that the information provided in this Expression of Interest is true and "
    "accurate to the best of my/our knowledge. Submission of this form does not guarantee acceptance "
    "or issuance of any insurance policy or partnership."
)


def _extract_doc_text(path: str) -> str:
    ext = Path(path).suffix.lower()

    if ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if ext == ".pdf":
        if PdfReader is None:
            raise RuntimeError("PDF reader library not installed. Install `pypdf`.")
        reader = PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in reader.pages)

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _clean_value(value: str | None) -> str:
    if not value:
        return ""
    v = re.sub(r"\s+", " ", value).strip()
    v = v.strip("-_:. ")
    return "" if v.lower() in {"na", "n/a", "not available", "nil", "none"} else v


def _sanitize_field_value(key: str, value: str | None) -> str:
    v = _clean_value(value)
    if not v:
        return ""

    if len(v) > FIELD_MAX_LEN.get(key, 120):
        return ""

    low = v.lower()
    if any(token in low for token in NOISE_TOKENS):
        return ""

    if key == "email_address":
        return v if re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", v) else ""

    if key == "phone_number":
        digits = re.sub(r"\D", "", v)
        return v if len(digits) >= 7 else ""

    if key == "website":
        has_domain = re.search(r"(https?://|www\.|[A-Za-z0-9-]+\.[A-Za-z]{2,})", v)
        return v if has_domain else ""

    if key in {"city", "state", "country"}:
        if ":" in v or len(v.split()) > 6:
            return ""
        return v

    if key == "postal_code":
        return v if re.search(r"[A-Za-z0-9]", v) else ""

    if key == "expected_sum_insured":
        # Should usually contain currency/number-like content.
        return v if re.search(r"\d", v) else ""

    if key in {"pan_tax_id", "gst_number"}:
        if re.search(r"[.!?].+[.!?]", v):
            return ""
        return v

    return v


def _extract_single_field(text: str, label_pattern: str, stop_labels: list[str]) -> str:
    stop = "|".join(stop_labels)
    pattern = rf"{label_pattern}\s*[:\-]?\s*(.+?)(?=\n\s*(?:{stop})\s*[:\-]|$)"
    m = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
    if not m:
        # Inline single-line fallback
        m2 = re.search(rf"{label_pattern}\s*[:\-]?\s*([^\n]+)", text, flags=re.IGNORECASE)
        return _clean_value(m2.group(1)) if m2 else ""
    return _clean_value(m.group(1))


def extract_broker_fields(doc_text: str) -> dict[str, str]:
    stop_labels = [
        r"Registered Address", r"City", r"State", r"Country", r"Postal Code",
        r"Phone Number", r"Email Address", r"Website(?:\s*\(if applicable\))?",
        r"Expected Sum Insured(?:\s*/\s*Coverage Amount)?", r"PAN(?:\s*/\s*Tax ID\s*/\s*Business ID)?",
        r"GST Number(?:\s*\(if applicable\))?", r"Regulatory Licenses(?:\s*/\s*Certifications)?",
    ]

    fields = {
        "registered_address": _extract_single_field(doc_text, r"Registered Address", stop_labels),
        "city": _extract_single_field(doc_text, r"City", stop_labels),
        "state": _extract_single_field(doc_text, r"State", stop_labels),
        "country": _extract_single_field(doc_text, r"Country", stop_labels),
        "postal_code": _extract_single_field(doc_text, r"Postal Code", stop_labels),
        "phone_number": _extract_single_field(doc_text, r"Phone Number", stop_labels),
        "email_address": _extract_single_field(doc_text, r"Email Address", stop_labels),
        "website": _extract_single_field(doc_text, r"Website(?:\s*\(if applicable\))?", stop_labels),
        "expected_sum_insured": _extract_single_field(doc_text, r"Expected Sum Insured(?:\s*/\s*Coverage Amount)?", stop_labels),
        "pan_tax_id": _extract_single_field(doc_text, r"PAN(?:\s*/\s*Tax ID\s*/\s*Business ID)?", stop_labels),
        "gst_number": _extract_single_field(doc_text, r"GST Number(?:\s*\(if applicable\))?", stop_labels),
        "regulatory_licenses": _extract_single_field(doc_text, r"Regulatory Licenses(?:\s*/\s*Certifications)?", stop_labels),
        "claims_history": _extract_single_field(doc_text, r"Claims History(?:\s*\(last\s*3[–-]5\s*years\))?", stop_labels),
    }
    fields = {k: _sanitize_field_value(k, v) for k, v in fields.items()}

    # LLM fallback only for missing fields
    missing = [k for k, v in fields.items() if not v]
    if missing:
        prompt = f"""
Extract the following fields from this broker submission text and return ONLY JSON.
Missing fields to extract: {missing}
Keys allowed: {list(FIELD_LABELS.keys())}
Use empty string if not found.

TEXT:
{doc_text[:12000]}
"""
        llm_raw = call_llm(prompt)
        m = re.search(r"\{.*\}", llm_raw, flags=re.DOTALL)
        if m:
            try:
                parsed = json.loads(m.group(0))
                for k in missing:
                    fields[k] = _sanitize_field_value(k, str(parsed.get(k, "")))
            except Exception:
                pass

    return fields


def _extract_risk_profile_json(user_prompt: str, doc_text: str, broker_fields: dict[str, str]) -> dict[str, Any]:
    base = {
        "lob": _infer_lob_hint(f"{user_prompt}\n{doc_text}") or "Not specified",
        "tiv": broker_fields.get("expected_sum_insured") or "Not specified",
        "turnover": "Not specified",
        "sites": "Not specified",
    }
    if not doc_text.strip():
        return base

    prompt = f"""
Extract a concise risk profile JSON from the broker submission.
Return ONLY valid JSON with keys exactly:
- lob
- tiv
- turnover
- sites

If unavailable, return "Not specified".

USER PROMPT:
{user_prompt}

BROKER FIELDS:
{json.dumps(broker_fields, indent=2)}

DOCUMENT TEXT:
{doc_text[:12000]}
"""
    try:
        raw = call_llm(prompt)
        match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
        if not match:
            return base
        parsed = json.loads(match.group(0))
        out = {
            "lob": str(parsed.get("lob") or base["lob"]).strip(),
            "tiv": str(parsed.get("tiv") or base["tiv"]).strip(),
            "turnover": str(parsed.get("turnover") or base["turnover"]).strip(),
            "sites": str(parsed.get("sites") or base["sites"]).strip(),
        }
        return out
    except Exception:
        return base


def _clean_geo_candidate(value: str) -> str:
    token = re.sub(r"\s+", " ", str(value or "")).strip(" ,;:.-")
    token = re.sub(r"\b(?:and|or|the)\b$", "", token, flags=re.IGNORECASE).strip(" ,;:.-")
    if not token:
        return ""
    if len(token) < 2 or len(token) > 60:
        return ""
    # Avoid instruction/noise fragments leaking into geography output.
    noise_phrases = [
        "summarize the attached broker",
        "submission document",
        "international casualty",
        "manufacturing sites",
        "distribution hubs",
        "expected sum insured",
        "coverage amount",
        "registered address",
        "not specified",
    ]
    low = token.lower()
    if any(p in low for p in noise_phrases):
        return ""
    noise_words = {
        "summarize", "attached", "broker", "submission", "document", "casualty", "manufacturing",
        "distribution", "hub", "hubs", "site", "sites", "turnover", "insured", "coverage", "amount",
    }
    word_hits = sum(1 for w in re.findall(r"[A-Za-z]+", low) if w in noise_words)
    if word_hits >= 2:
        return ""
    if not re.search(r"[A-Za-z]", token):
        return ""
    return token


def _extract_geo_tokens(user_prompt: str, risk_profile: dict[str, Any] | None, broker_fields: dict[str, str] | None) -> list[str]:
    profile = risk_profile or {}
    fields = broker_fields or {}

    raw_values = [
        str(fields.get("city") or ""),
        str(fields.get("state") or ""),
        str(fields.get("country") or ""),
        str(profile.get("sites") or ""),
        str(fields.get("registered_address") or ""),
    ]

    tokens: list[str] = []
    seen: set[str] = set()

    for value in raw_values:
        if not value:
            continue
        parts = re.split(r"[\n,;/|]+", value)
        for part in parts:
            cleaned = _clean_geo_candidate(part)
            if not cleaned:
                continue
            key = cleaned.lower()
            if key in seen:
                continue
            seen.add(key)
            tokens.append(cleaned)

    # Minimal fallback from prompt only when no structured location was extracted.
    if not tokens and user_prompt:
        prompt_candidates = re.findall(r"\b(?:in|at|across|from)\s+([A-Za-z][A-Za-z\s]{2,30})", user_prompt, flags=re.IGNORECASE)
        for candidate in prompt_candidates:
            cleaned = _clean_geo_candidate(candidate)
            if not cleaned:
                continue
            key = cleaned.lower()
            if key in seen:
                continue
            seen.add(key)
            tokens.append(cleaned)

    return tokens[:8]


def _ensure_sentence(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if not cleaned:
        return ""
    if cleaned.endswith(("...", ".", "!", "?")):
        return cleaned
    return f"{cleaned}."


def _normalize_complete_sentences(text: str, max_sentences: int = 5) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if not cleaned:
        return ""
    parts = re.split(r"(?<=[.!?])\s+", cleaned)
    if len(parts) == 1:
        # Fallback split for fragment-style outputs.
        parts = [p.strip() for p in re.split(r"\s*[;|]\s*", cleaned) if p.strip()]
    sentences = [_ensure_sentence(p) for p in parts if p.strip()]
    if not sentences:
        return _ensure_sentence(cleaned)
    return " ".join(sentences[:max_sentences])


def _looks_fragmented_geo_summary(text: str) -> bool:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if not cleaned:
        return True
    # Heading-like snippets such as "4. Contents" or "3. The Role of ...".
    if re.match(r"^\d+\.?\s+[A-Za-z][A-Za-z\s:&()/-]{1,90}$", cleaned):
        return True
    if re.match(r"^(contents|table of contents|introduction|overview)$", cleaned, flags=re.IGNORECASE):
        return True
    words = re.findall(r"[A-Za-z]+", cleaned)
    if len(words) < 8:
        return True
    return False


def _llm_geo_link_summary(title: str, summary: str) -> str:
    prompt = f"""
You are an insurance underwriting analyst.

Rewrite the following web snippet into exactly 2 to 3 complete sentences.
Requirements:
- Complete sentences only (no fragments, no headings).
- Focus on geography-based risk signals relevant to calamity, crime, and strike risk.
- Keep it concise and professional for underwriting review.
- Do not include bullet points.

Title: {title}
Snippet: {summary}
"""
    try:
        raw = call_llm(prompt).strip()
        return _normalize_complete_sentences(raw, max_sentences=3)
    except Exception:
        return ""


def _build_geo_link_summary(title: str, summary: str) -> str:
    normalized = _normalize_complete_sentences(summary, max_sentences=3)
    sentences = [s for s in re.split(r"(?<=[.!?])\s+", normalized) if s.strip()]
    if len(sentences) >= 2 and not _looks_fragmented_geo_summary(normalized):
        llm_out = _llm_geo_link_summary(title, normalized)
        if llm_out:
            return llm_out
        return " ".join(sentences[:3])

    title_clean = re.sub(r"\s+", " ", str(title or "")).strip()
    if not title_clean:
        title_clean = "Referenced external source"
    title_clean = re.sub(r"^[0-9]+\.\s*", "", title_clean).strip()

    llm_out = _llm_geo_link_summary(title_clean, normalized or summary)
    if llm_out:
        llm_sentences = [s for s in re.split(r"(?<=[.!?])\s+", llm_out) if s.strip()]
        if len(llm_sentences) >= 2:
            return " ".join(llm_sentences[:3])

    fallback_extra = [
        _ensure_sentence(f"The source '{title_clean}' provides geography-related external risk context for underwriting review"),
        "This signal is considered for calamity, crime, and strike exposure at the insured geography.",
        "The insight is used as supporting context in the geo-risk scoring layer and does not replace internal policy controls.",
    ]
    base = sentences[0] if sentences and not _looks_fragmented_geo_summary(sentences[0]) else ""
    parts = [_ensure_sentence(base)] if base else []
    parts.extend(fallback_extra)
    return " ".join([p for p in parts if p])


def _compute_web_geo_risk(
    user_prompt: str,
    web_summary: str,
    web_links: list | None,
    risk_profile: dict[str, Any] | None,
    broker_fields: dict[str, str] | None,
) -> dict[str, Any]:
    link_text = " ".join(
        [
            f"{item[0]} {item[1]}"
            for item in (web_links or [])
            if isinstance(item, (list, tuple)) and len(item) >= 2
        ]
    )
    low = f"{(web_summary or '')} {link_text}".lower()
    geo_tokens = _extract_geo_tokens(user_prompt, risk_profile, broker_fields)
    if "global" in (user_prompt or "").lower() or "international" in (user_prompt or "").lower():
        if not any(str(t).strip().lower() == "global" for t in geo_tokens):
            geo_tokens.append("Global")

    severe_markers = [
        "severe", "major", "critical", "high", "escalation", "active", "widespread", "persistent"
    ]
    moderate_markers = [
        "moderate", "watch", "elevated", "potential", "localized", "sporadic"
    ]

    hazard_groups = {
        "calamity": ["flood", "earthquake", "cyclone", "storm", "wildfire", "drought", "hurricane", "landslide", "natcat", "catastrophe"],
        "crime": ["crime", "theft", "burglary", "vandalism", "riot", "violence", "terror", "kidnap", "fraud", "cyber attack"],
        "strike": ["strike", "labor unrest", "industrial action", "walkout", "protest", "civil commotion", "srcc"],
    }

    score = 15.0
    drivers: list[str] = []
    detected: dict[str, bool] = {k: False for k in hazard_groups}

    for group, words in hazard_groups.items():
        group_hits = [w for w in words if w in low]
        if not group_hits:
            continue
        detected[group] = True

        points = 10.0
        points += min(12.0, 2.5 * (len(group_hits) - 1))
        if any(m in low for m in severe_markers):
            points += 12.0
        elif any(m in low for m in moderate_markers):
            points += 6.0

        score += points
        drivers.append(_ensure_sentence(f"{group.title()} risk indicators were detected: {', '.join(group_hits[:3])}"))

    if geo_tokens:
        score += min(12.0, len(geo_tokens) * 1.2)
        drivers.append(_ensure_sentence(f"Geography analyzed for hazard context: {', '.join(geo_tokens[:4])}"))

    score = max(0.0, min(100.0, round(score, 2)))
    if score < WEB_RISK_BANDS["low_max"]:
        level = "LOW"
    elif score <= WEB_RISK_BANDS["moderate_max"]:
        level = "MODERATE"
    else:
        level = "HIGH"

    if not drivers:
        drivers = ["No significant calamity, crime, or strike indicators were found in the web summary."]

    return {
        "score": score,
        "level": level,
        "detected_hazards": detected,
        "geo_tokens": geo_tokens,
        "drivers": [_ensure_sentence(d) for d in drivers],
    }


def summarize_doc_with_instruction(doc_path: str, instruction: str) -> str:
    if not doc_path or not os.path.exists(doc_path):
        return "No document uploaded."

    try:
        content = _extract_doc_text(doc_path)
    except Exception as exc:
        return f"Failed to read document: {exc}"

    prompt = f"""
You are an underwriting analyst.

Instruction:
{instruction}

Submission document:
{content[:12000]}

Provide:
1) A concise summary of the submission.
2) Key underwriting pros and cons.
3) A clear recommendation on whether to proceed, and why.
"""
    return call_llm(prompt).strip()


def build_vanna_prompt(user_prompt: str) -> str:
    reserve_classes = [
        "aviation",
        "property",
        "casualty",
        "marine",
        "motor",
        "energy",
        "financial lines",
    ]

    prompt = user_prompt.lower()
    matched = next((rc for rc in reserve_classes if rc in prompt), None)

    if not matched:
        return "Show ultimate loss for the last 10 years"

    return f"Show ultimate loss for {matched} for the last 10 years"


def build_serp_prompt(user_prompt: str) -> str:
    """
    Build a focused external-search prompt instead of passing the full user text.
    Keeps SERP intent narrow: market/industry insights, not document summarization noise.
    """
    prompt = (user_prompt or "").lower()

    reserve_classes = [
        "aviation",
        "property",
        "casualty",
        "marine",
        "motor",
        "energy",
        "financial lines",
        "construction",
    ]
    matched = next((rc for rc in reserve_classes if rc in prompt), "casualty")

    geography = "global"
    if "international" in prompt:
        geography = "international"
    elif "us" in prompt or "united states" in prompt or "north america" in prompt:
        geography = "north america"
    elif "europe" in prompt:
        geography = "europe"
    elif "asia" in prompt:
        geography = "asia"

    metric = "ultimate losses"
    if "loss ratio" in prompt:
        metric = "loss ratios"
    elif "incurred" in prompt:
        metric = "incurred and ultimate losses"

    return (
        f"Fetch latest {geography} market web insights for {matched} lines, "
        f"with focus on {metric}, pricing trend, claims trend, and underwriting appetite."
    )


def _resolve_eoi_db_path() -> Path:
    candidates = [
        os.getenv("EOI_DB_PATH"),
        os.getenv("UNDERWRITER_DB_PATH"),
        str(DEFAULT_EOI_DB_PATH),
        str(Path(DB_PATH)),
    ]
    for candidate in candidates:
        if not candidate:
            continue
        p = Path(candidate).expanduser()
        if not p.is_absolute():
            p = PROJECT_ROOT / p
        p = p.resolve()
        if p.exists():
            return p
    return DEFAULT_EOI_DB_PATH


def _resolve_underwriting_table(conn: sqlite3.Connection) -> str:
    rows = conn.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name").fetchall()
    names = [str(r[0]) for r in rows if r and r[0]]
    if not names:
        raise RuntimeError("No tables found in underwriting database.")

    preferred = [
        "underwriting_dataset",
        "underwriter_data",
        "underwriting_data",
        "portfolio",
    ]
    lowered = {n.lower(): n for n in names}
    for key in preferred:
        if key in lowered:
            return lowered[key]

    return names[0]


def _to_float(value: Any) -> float | None:
    try:
        if value is None:
            return None
        if isinstance(value, (int, float)):
            return float(value)
        txt = re.sub(r"[^0-9.\-]", "", str(value))
        if not txt:
            return None
        return float(txt)
    except Exception:
        return None


AMOUNT_SCALE_FACTORS = {
    "k": 1_000.0,
    "thousand": 1_000.0,
    "m": 1_000_000.0,
    "mn": 1_000_000.0,
    "mm": 1_000_000.0,
    "million": 1_000_000.0,
    "b": 1_000_000_000.0,
    "bn": 1_000_000_000.0,
    "billion": 1_000_000_000.0,
    "lakh": 100_000.0,
    "lac": 100_000.0,
    "crore": 10_000_000.0,
    "cr": 10_000_000.0,
}


def _parse_amount_with_scale(text: str | None) -> float | None:
    s = str(text or "").strip().lower()
    if not s:
        return None
    m = re.search(
        r"([-+]?\d[\d,]*(?:\.\d+)?)\s*(k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr)?\b",
        s,
        flags=re.IGNORECASE,
    )
    if not m:
        return None
    try:
        base = float(m.group(1).replace(",", ""))
    except Exception:
        return None
    scale = (m.group(2) or "").lower()
    factor = AMOUNT_SCALE_FACTORS.get(scale, 1.0)
    return base * factor


def _extract_yearly_incurred_amounts(text: str) -> list[float]:
    values: list[float] = []
    # Examples: "2020: 1.2 million", "FY2021 - 850000"
    year_line = re.findall(
        r"(?:fy\s*)?(?:19|20)\d{2}\s*[:\-]\s*([$€£]?\s*[-+]?\d[\d,]*(?:\.\d+)?(?:\s*(?:k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr))?)",
        text,
        flags=re.IGNORECASE,
    )
    for raw in year_line:
        amount = _parse_amount_with_scale(raw)
        if amount is not None:
            values.append(amount)
    return values


def _extract_year_span(text: str, default_years: float = 5.0) -> float:
    match = re.search(r"(?:over|for|last|across|during)\s*([0-9]{1,2})\s*years?", text, flags=re.IGNORECASE)
    years = _to_float(match.group(1)) if match else default_years
    return years if years and years > 0 else default_years


def _extract_yearly_claim_counts(text: str) -> list[float]:
    values: list[float] = []
    # Examples: "2021: 12 claims", "FY2022 - 8"
    matches = re.findall(
        r"(?:fy\s*)?(?:19|20)\d{2}\s*[:\-]\s*([0-9]+(?:\.[0-9]+)?)\s*(?:claims?)?",
        text,
        flags=re.IGNORECASE,
    )
    for raw in matches:
        val = _to_float(raw)
        if val is not None:
            values.append(val)
    return values


def _extract_yearly_largest_losses(text: str) -> list[float]:
    values: list[float] = []

    patterns = [
        # Example: "FY2022: Largest Loss 1.25 million"
        r"(?:fy\s*)?(?:19|20)\d{2}\s*[:\-]\s*(?:largest\s*(?:single\s*)?loss|max(?:imum)?\s*loss)[^0-9]{0,25}([$€£]?\s*[-+]?\d[\d,]*(?:\.\d+)?(?:\s*(?:k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr))?)",
        # Example: "Largest loss ... FY2022 ... 1.25m"
        r"(?:largest\s*(?:single\s*)?loss|max(?:imum)?\s*loss)[^\n\r]{0,60}(?:fy\s*)?(?:19|20)\d{2}[^\n\r]{0,20}([$€£]?\s*[-+]?\d[\d,]*(?:\.\d+)?(?:\s*(?:k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr))?)",
    ]

    for pat in patterns:
        matches = re.findall(pat, text, flags=re.IGNORECASE)
        for raw in matches:
            amount = _parse_amount_with_scale(raw)
            if amount is not None:
                values.append(amount)

    return values


def _extract_loss_history_rows(text: str) -> list[tuple[int, float, float]]:
    """
    Extract tabular loss-history rows in the common format:
    YEAR  NUMBER_OF_CLAIMS  INCURRED  LARGEST_LOSS
    """
    rows: list[tuple[int, float, float]] = []
    pattern = re.compile(
        r"\b((?:19|20)\d{2})\b\s+([0-9]+(?:\.[0-9]+)?)\s+([0-9][0-9,]*(?:\.[0-9]+)?)\s+([0-9][0-9,]*(?:\.[0-9]+)?)",
        flags=re.IGNORECASE,
    )
    for m in pattern.finditer(text):
        year = int(m.group(1))
        if year < 1990 or year > 2100:
            continue
        incurred = _to_float(m.group(3))
        largest = _to_float(m.group(4))
        if incurred is None or largest is None:
            continue
        rows.append((year, incurred, largest))
    return rows


def _is_plausible_loss_amount(value: float | None) -> bool:
    if value is None:
        return False
    # Guard against capturing section numbers like "5." as severity/loss.
    return value >= 1_000.0


def _parse_ratio_from_text(value: str | None) -> float | None:
    s = str(value or "").strip().lower()
    if not s:
        return None
    m = re.search(r"([-+]?\d[\d,]*(?:\.\d+)?)\s*(%|percent|percentage)?\b", s, flags=re.IGNORECASE)
    if not m:
        return None
    try:
        num = float(m.group(1).replace(",", ""))
    except Exception:
        return None
    # Keep as percent-scale value when percentage marker exists.
    if m.group(2):
        return num
    return num


def _sql_literal(value: Any) -> str:
    if value is None:
        return "NULL"
    if isinstance(value, bool):
        return "1" if value else "0"
    if isinstance(value, (int, float)):
        return str(value)
    text = str(value).replace("'", "''")
    return f"'{text}'"


def _render_sql_with_params(sql: str, params: list[Any]) -> str:
    rendered = str(sql or "")
    for p in params:
        rendered = rendered.replace("?", _sql_literal(p), 1)
    return rendered


def _extract_client_loss_ratio(
    user_prompt: str,
    broker_fields: dict[str, str] | None,
    broker_text: str | None,
) -> tuple[float | None, str]:
    """
    Try to extract client-specific loss ratio from available text sources.
    Returns (value, source_label). Value can be in either ratio or percent scale.
    """
    fields = broker_fields or {}
    candidates = [
        ("claims_history", str(fields.get("claims_history") or "")),
        ("user_prompt", str(user_prompt or "")),
        ("broker_text", str(broker_text or "")),
    ]
    patterns = [
        r"loss\s*ratio[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*%",
        r"loss\s*ratio[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*(?:percent|percentage)",
        r"loss\s*ratio[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)",
    ]
    for src, text in candidates:
        low = text.lower()
        for pat in patterns:
            m = re.search(pat, low, flags=re.IGNORECASE)
            if not m:
                continue
            val = _parse_ratio_from_text(m.group(1))
            if val is not None:
                return val, src
    return None, "not_found"


def _extract_client_frequency(
    user_prompt: str,
    broker_fields: dict[str, str] | None,
    broker_text: str | None,
) -> tuple[float | None, str]:
    fields = broker_fields or {}
    candidates = [
        ("claims_history", str(fields.get("claims_history") or "")),
        ("user_prompt", str(user_prompt or "")),
        ("broker_text", str(broker_text or "")),
    ]
    for src, text in candidates:
        low = text.lower()

        yearly_counts = _extract_yearly_claim_counts(low)
        if len(yearly_counts) >= 1:
            return float(sum(yearly_counts) / len(yearly_counts)), src

        # Primary signal: "Number of Claims" entries in broker submission.
        num_claims_patterns = [
            r"number\s*of\s*claims?[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*(?:per\s*year|/yr|yearly)?",
            r"claims?\s*count[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*(?:per\s*year|/yr|yearly)?",
            r"([0-9]+(?:\.[0-9]+)?)\s*claims?\s*(?:per\s*year|/yr|yearly)",
        ]
        for pat in num_claims_patterns:
            m = re.search(pat, low, flags=re.IGNORECASE)
            if not m:
                continue
            val = _to_float(m.group(1))
            if val is not None:
                return val, src

        total_claims_match = re.search(
            r"([0-9]+(?:\.[0-9]+)?)\s*claims?\s*(?:over|for|across|during|in)\s*([0-9]{1,2})\s*years?",
            low,
            flags=re.IGNORECASE,
        )
        if total_claims_match:
            total_claims = _to_float(total_claims_match.group(1))
            years = _to_float(total_claims_match.group(2))
            if total_claims is not None and years and years > 0:
                return total_claims / years, src

        # If claims count is stated without per-year context, normalize by detected year span.
        generic_claims = re.search(r"(?:total\s*)?claims?[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)", low, flags=re.IGNORECASE)
        if generic_claims:
            total_claims = _to_float(generic_claims.group(1))
            if total_claims is not None:
                years = _extract_year_span(low, default_years=5.0)
                return total_claims / years, src

        # Fallback only: if explicit claims counts are absent, use reported frequency value.
        freq_fallback = re.search(r"(?:claims?\s*frequency|frequency)[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)", low, flags=re.IGNORECASE)
        if freq_fallback:
            val = _to_float(freq_fallback.group(1))
            if val is not None:
                return val, src
    return None, "not_found"


def _extract_client_severity(
    user_prompt: str,
    broker_fields: dict[str, str] | None,
    broker_text: str | None,
) -> tuple[float | None, str]:
    fields = broker_fields or {}
    candidates = [
        ("broker_text", str(broker_text or "")),
        ("claims_history", str(fields.get("claims_history") or "")),
        ("user_prompt", str(user_prompt or "")),
    ]
    patterns = [
        r"(?:largest|max(?:imum)?)\s*(?:single\s*)?loss[^0-9]{0,25}([$€£]?\s*[-+]?\d[\d,]*(?:\.\d+)?(?:\s*(?:k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr))?)",
    ]
    for src, text in candidates:
        low = text.lower()
        # Highest-priority path: table rows from loss history section.
        table_rows = _extract_loss_history_rows(low)
        if table_rows:
            losses = [r[2] for r in table_rows if _is_plausible_loss_amount(r[2])]
            if losses:
                return float(sum(losses) / len(losses)), src

        yearly_losses = _extract_yearly_largest_losses(low)
        yearly_losses = [v for v in yearly_losses if _is_plausible_loss_amount(v)]
        if len(yearly_losses) >= 1:
            # Use average annual largest loss to keep it on per-year basis.
            return float(sum(yearly_losses) / len(yearly_losses)), src
        for pat in patterns:
            m = re.search(pat, low, flags=re.IGNORECASE)
            if not m:
                continue
            val = _parse_amount_with_scale(m.group(1))
            if _is_plausible_loss_amount(val):
                return val, src
    return None, "not_found"


def _extract_client_incurred_per_year(
    user_prompt: str,
    broker_fields: dict[str, str] | None,
    broker_text: str | None,
) -> tuple[float | None, str]:
    fields = broker_fields or {}
    candidates = [
        ("claims_history", str(fields.get("claims_history") or "")),
        ("user_prompt", str(user_prompt or "")),
        ("broker_text", str(broker_text or "")),
    ]
    for src, text in candidates:
        low = text.lower()
        per_year_patterns = [
            r"(?:incurred(?:\s*loss)?(?:\s*per\s*year)?|annual\s*incurred)[^0-9]{0,25}([$€£]?\s*[-+]?\d[\d,]*(?:\.\d+)?(?:\s*(?:k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr))?)",
        ]
        for pat in per_year_patterns:
            m = re.search(pat, low, flags=re.IGNORECASE)
            if not m:
                continue
            val = _parse_amount_with_scale(m.group(1))
            if val is not None:
                return val, src

        yearly_values = _extract_yearly_incurred_amounts(low)
        if len(yearly_values) >= 2:
            return float(sum(yearly_values) / len(yearly_values)), src

        total_patterns = [
            r"(?:total\s*)?incurred(?:\s*loss)?[^0-9]{0,25}([$€£]?\s*[-+]?\d[\d,]*(?:\.\d+)?(?:\s*(?:k|thousand|m|mn|mm|million|b|bn|billion|lakh|lac|crore|cr))?)",
        ]
        years = _extract_year_span(low, default_years=5.0)
        for pat in total_patterns:
            m = re.search(pat, low, flags=re.IGNORECASE)
            if not m:
                continue
            total_val = _parse_amount_with_scale(m.group(1))
            if total_val is not None:
                return total_val / years, src
    return None, "not_found"


def _extract_class_of_business(
    user_prompt: str,
    broker_text: str | None,
) -> tuple[str | None, str]:
    text = str(broker_text or "")
    if text:
        patterns = [
            r"class\s*of\s*business\s*[:\-]\s*([^\n\r]{2,120})",
            r"line\s*of\s*business\s*[:\-]\s*([^\n\r]{2,120})",
            r"lob\s*[:\-]\s*([^\n\r]{2,120})",
        ]
        for pat in patterns:
            m = re.search(pat, text, flags=re.IGNORECASE)
            if not m:
                continue
            raw = _clean_value(m.group(1))
            if not raw:
                continue
            # Keep first token-like phrase to avoid carrying trailing narrative text.
            candidate = re.split(r"[.;|]", raw)[0].strip()
            if candidate:
                return candidate, "broker_submission"

    inferred_from_doc = _infer_lob_hint(text)
    if inferred_from_doc:
        return inferred_from_doc, "broker_submission_inferred"

    inferred_from_prompt = _infer_lob_hint(user_prompt or "")
    if inferred_from_prompt:
        return inferred_from_prompt, "user_prompt_inferred"

    return None, "not_found"


def _align_metric_scale(client_value: float, portfolio_series: pd.Series, ratio_mode: bool = False) -> float:
    """
    Align client scale to portfolio scale:
    - ratio_mode=True handles ratio fields where source may be in percent or fraction
    - otherwise return numeric value unchanged
    """
    if portfolio_series.empty:
        return client_value
    if not ratio_mode:
        return client_value
    med = float(portfolio_series.median())
    portfolio_is_percent = med > 1.5
    client_is_percent = client_value > 1.5
    if portfolio_is_percent and not client_is_percent:
        return client_value * 100.0
    if not portfolio_is_percent and client_is_percent:
        return client_value / 100.0
    return client_value


def _run_internal_sql(
    vanna_prompt: str,
    user_prompt: str = "",
    broker_fields: dict[str, str] | None = None,
    broker_text: str | None = None,
) -> tuple[str | None, pd.DataFrame | None]:
    broker_fields = broker_fields or {}
    extracted_cob, _ = _extract_class_of_business(user_prompt=user_prompt, broker_text=broker_text)
    fallback_lob_hint = _infer_lob_hint(f"{user_prompt} {broker_fields.get('claims_history', '')}") or ""
    lob_hint = extracted_cob or fallback_lob_hint
    db_path = _resolve_eoi_db_path()
    sql_query = ""
    sql_query_display = ""

    try:
        conn = sqlite3.connect(str(db_path))
        try:
            table = _resolve_underwriting_table(conn)
            col_rows = conn.execute(f"PRAGMA table_info('{table}')").fetchall()
            columns = {str(r[1]).lower() for r in col_rows if len(r) > 1}

            def has_col(name: str) -> bool:
                return name.lower() in columns

            def agg_or_null(col: str, func: str, alias: str) -> str:
                return f"{func}({col}) AS {alias}" if has_col(col) else f"NULL AS {alias}"

            select_parts = [
                "COUNT(*) AS account_count",
                agg_or_null("loss_ratio", "AVG", "avg_loss_ratio"),
                agg_or_null("claims_frequency", "AVG", "avg_claims_frequency"),
                agg_or_null("largest_single_loss", "AVG", "avg_largest_single_loss"),
                agg_or_null("ultimate_premium", "AVG", "avg_ultimate_premium"),
                agg_or_null("incurred_loss", "AVG", "avg_incurred_loss"),
                agg_or_null("incurred_loss", "SUM", "total_incurred_loss"),
            ]

            filters = []
            params: list[Any] = []
            if lob_hint:
                if has_col("class_of_business"):
                    lob_filter_parts = []
                    lob_filter_parts.append("LOWER(class_of_business) LIKE ?")
                    params.append(f"%{lob_hint.lower()}%")
                    filters.append("(" + " OR ".join(lob_filter_parts) + ")")

            sql_query = f"SELECT {', '.join(select_parts)} FROM {table}"
            if filters:
                sql_query += " WHERE " + " AND ".join(filters)
            sql_query_display = _render_sql_with_params(sql_query, params)
            stats_df = pd.read_sql_query(sql_query, conn, params=params)

            def _load_distribution(metric_col: str) -> pd.Series:
                if not has_col(metric_col):
                    return pd.Series(dtype="float64")
                dist_query = f"SELECT {metric_col} FROM {table} WHERE {metric_col} IS NOT NULL"
                dist_params: list[Any] = []
                if filters:
                    dist_query += " AND " + " AND ".join(filters)
                    dist_params = list(params)
                dist_df = pd.read_sql_query(dist_query, conn, params=dist_params)
                return pd.to_numeric(dist_df[metric_col], errors="coerce").dropna()

            loss_ratio_series = _load_distribution("loss_ratio")
            claims_freq_series = _load_distribution("claims_frequency")
            severity_series = _load_distribution("largest_single_loss")
            incurred_series = _load_distribution("incurred_loss")
        finally:
            conn.close()

        if stats_df.empty:
            return sql_query_display or sql_query, pd.DataFrame([{"Error": "No rows returned from underwriting_dataset."}])

        client_loss_ratio_raw, client_lr_source = _extract_client_loss_ratio(
            user_prompt=user_prompt,
            broker_fields=broker_fields,
            broker_text=broker_text,
        )
        client_freq_raw, client_freq_source = _extract_client_frequency(
            user_prompt=user_prompt,
            broker_fields=broker_fields,
            broker_text=broker_text,
        )
        client_severity_raw, client_severity_source = _extract_client_severity(
            user_prompt=user_prompt,
            broker_fields=broker_fields,
            broker_text=broker_text,
        )
        client_incurred_raw, client_incurred_source = _extract_client_incurred_per_year(
            user_prompt=user_prompt,
            broker_fields=broker_fields,
            broker_text=broker_text,
        )

        percentile = 50.0
        percentile_basis = "neutral_default"
        client_loss_ratio_aligned = None
        if not loss_ratio_series.empty and client_loss_ratio_raw is not None:
            client_loss_ratio_aligned = _align_metric_scale(client_loss_ratio_raw, loss_ratio_series, ratio_mode=True)
            percentile = float((loss_ratio_series <= client_loss_ratio_aligned).mean() * 100.0)
            percentile_basis = "client_loss_ratio"

        claims_freq_percentile = 50.0
        client_freq_aligned = client_freq_raw
        if not claims_freq_series.empty and client_freq_raw is not None:
            client_freq_aligned = _align_metric_scale(client_freq_raw, claims_freq_series)
            claims_freq_percentile = float((claims_freq_series <= client_freq_aligned).mean() * 100.0)

        severity_percentile = 50.0
        client_severity_aligned = client_severity_raw
        if not severity_series.empty and client_severity_raw is not None:
            client_severity_aligned = _align_metric_scale(client_severity_raw, severity_series)
            severity_percentile = float((severity_series <= client_severity_aligned).mean() * 100.0)

        incurred_percentile = 50.0
        client_incurred_aligned = client_incurred_raw
        if not incurred_series.empty and client_incurred_raw is not None:
            client_incurred_aligned = _align_metric_scale(client_incurred_raw, incurred_series)
            incurred_percentile = float((incurred_series <= client_incurred_aligned).mean() * 100.0)

        stats_df["loss_ratio_percentile"] = percentile
        stats_df["loss_ratio_percentile_basis"] = percentile_basis
        stats_df["client_loss_ratio"] = client_loss_ratio_aligned
        stats_df["client_loss_ratio_source"] = client_lr_source
        stats_df["claims_frequency_percentile"] = claims_freq_percentile
        stats_df["client_claims_frequency"] = client_freq_aligned
        stats_df["client_claims_frequency_source"] = client_freq_source
        stats_df["severity_percentile"] = severity_percentile
        stats_df["client_severity"] = client_severity_aligned
        stats_df["client_severity_source"] = client_severity_source
        stats_df["incurred_percentile"] = incurred_percentile
        stats_df["client_incurred_per_year"] = client_incurred_aligned
        stats_df["client_incurred_source"] = client_incurred_source
        stats_df["lob_hint"] = lob_hint or "Not inferred"
        stats_df["source_db"] = str(db_path)
        return sql_query_display or sql_query, stats_df
    except Exception as exc:
        # Keep backward compatibility by trying Vanna if direct DB benchmark fails.
        try:
            vanna_out = vanna_node({"user_prompt": vanna_prompt})
            vanna_query = vanna_out.get("sql_query") or sql_query
            result = vanna_out.get("sql_result")
            if isinstance(result, pd.DataFrame):
                return vanna_query, result
            if isinstance(result, list):
                return vanna_query, pd.DataFrame(result)
            if result is None:
                return vanna_query, pd.DataFrame()
            return vanna_query, pd.DataFrame([{"Result": str(result)}])
        except Exception:
            return sql_query, pd.DataFrame([{"Error": f"SQL Execution failed: {exc}"}])


def _normalize_loss_ratio(value: float | None) -> float:
    if value is None:
        return 55.0
    if value <= 40:
        return 15.0
    if value <= 60:
        return 30.0
    if value <= 80:
        return 50.0
    if value <= 100:
        return 70.0
    if value <= 120:
        return 85.0
    return 100.0


def _keyword_score(text: str, positive: list[str], negative: list[str], neutral_default: float = 50.0) -> float:
    low = (text or "").lower()
    pos = sum(1 for w in positive if w in low)
    neg = sum(1 for w in negative if w in low)
    if pos == 0 and neg == 0:
        return neutral_default
    if pos >= neg:
        return max(5.0, 45.0 - 7.0 * (pos - neg))
    return min(100.0, 55.0 + 9.0 * (neg - pos))


def _extract_conditions(doc_insights: str, intranet_summary: str, web_summary: str, risk_score: float) -> list[str]:
    conditions = [
        "Final underwriting approval subject to complete KYC and sanctions screening.",
        "Subject to policy wording alignment with internal underwriting guidelines.",
    ]
    low = f"{doc_insights}\n{intranet_summary}\n{web_summary}".lower()
    if "flood" in low or "earthquake" in low or "catastrophe" in low:
        conditions.append("Apply catastrophe sub-limit and consider peril-specific deductible.")
    if "litigation" in low or "adverse media" in low:
        conditions.append("Require enhanced due diligence and legal review before bind.")
    if "high" in low and "loss" in low:
        conditions.append("Apply claims loading and require risk-improvement plan.")
    if risk_score > SCORE_BANDS["write"]:
        conditions.append("Refer final terms and authority sign-off to senior underwriter.")
    return conditions[:5]


def _compute_risk_decision(
    user_prompt: str,
    broker_text: str,
    broker_fields: dict[str, str],
    risk_profile: dict[str, Any] | None,
    sql_result: pd.DataFrame | None,
    web_summary: str,
    web_risk: dict[str, Any] | None,
    intranet_summary: str,
    doc_insights: str,
) -> dict[str, Any]:
    row = {}
    if isinstance(sql_result, pd.DataFrame) and not sql_result.empty:
        row = sql_result.fillna("").iloc[0].to_dict()

    client_lr = _to_float(row.get("client_loss_ratio"))
    avg_loss_ratio = _to_float(row.get("avg_loss_ratio"))
    percentile = _to_float(row.get("loss_ratio_percentile"))
    claims_freq_percentile = _to_float(row.get("claims_frequency_percentile"))
    severity_percentile = _to_float(row.get("severity_percentile"))
    incurred_percentile = _to_float(row.get("incurred_percentile"))
    sum_insured = _to_float(broker_fields.get("expected_sum_insured"))
    profile = risk_profile or {}
    turnover_val = _to_float(profile.get("turnover"))
    sites_val = _to_float(profile.get("sites"))

    low_intranet = (intranet_summary or "").lower()
    low_web = (web_summary or "").lower()
    low_doc = (doc_insights or "").lower()
    combined = f"{user_prompt}\n{broker_text}\n{intranet_summary}\n{web_summary}"
    countries = set(re.findall(r"\b(?:in|across|from)\s+([A-Za-z ]{3,30})", combined, flags=re.IGNORECASE))
    web_risk_score = float((web_risk or {}).get("score") or 50.0)

    hard_rules = []
    if any(x in low_web for x in ["sanction", "ofac", "blacklist"]):
        hard_rules.append("Sanctions hit detected in external screening.")
    ratio_for_hard_rule = client_lr if client_lr is not None else avg_loss_ratio
    if ratio_for_hard_rule is not None:
        ratio_percent_value = ratio_for_hard_rule * 100.0 if ratio_for_hard_rule <= 1.5 else ratio_for_hard_rule
        if ratio_percent_value > 120:
            hard_rules.append(f"Average loss ratio above decline threshold ({ratio_percent_value:.1f} > 120).")
    if any(x in low_intranet for x in ["exceeds authority", "beyond authority", "referral required"]):
        hard_rules.append("Requested limit exceeds delegated authority.")
    if any(x in low_intranet for x in ["do not underwrite", "decline", "prohibited"]):
        hard_rules.append("Guideline hard-stop identified in intranet policy.")

    loss_quality_composite = max(0.0, min(100.0, percentile if percentile is not None else 50.0))
    loss_pattern_risk = (
        (claims_freq_percentile if claims_freq_percentile is not None else 50.0) * LOSS_PATTERN_WEIGHTS["claims_frequency_percentile"]
        + (severity_percentile if severity_percentile is not None else 50.0) * LOSS_PATTERN_WEIGHTS["severity_percentile"]
        + (incurred_percentile if incurred_percentile is not None else 50.0) * LOSS_PATTERN_WEIGHTS["incurred_percentile"]
    )

    metric_scores = {
        "loss_quality_composite": round(loss_quality_composite, 2),
        "loss_pattern_risk": round(max(0.0, min(100.0, loss_pattern_risk)), 2),
        "revenue_scale_risk": (
            70.0
            if (turnover_val or 0) >= 500000000
            else 60.0
            if (sum_insured or 0) >= 100000000
            else 40.0
            if (turnover_val or sum_insured)
            else 50.0
        ),
        "geographic_spread_risk": min(
            95.0,
            20.0 + 15.0 * max(0, int(sites_val) - 1) if sites_val is not None else 20.0 + 15.0 * max(0, len(countries) - 1),
        ),
        "risk_management_quality": 100.0 - _keyword_score(
            low_doc,
            positive=["strong controls", "certified", "iso", "compliant", "risk management framework"],
            negative=["weak control", "control gap", "incident", "non-compliance", "poor governance"],
            neutral_default=50.0,
        ),
        "external_risk": _keyword_score(
            low_web,
            positive=["stable", "favorable", "low volatility", "no major incident"],
            negative=["litigation", "adverse media", "catastrophe", "esg concern", "geopolitical"],
            neutral_default=web_risk_score,
        ),
        "coverage_complexity": _keyword_score(
            (user_prompt + " " + broker_text).lower(),
            positive=["standard coverage", "single line"],
            negative=["umbrella", "excess", "multi-country", "manuscript wording", "bespoke"],
            neutral_default=45.0,
        ),
        "guideline_fit": 100.0 - _keyword_score(
            low_intranet,
            positive=["within appetite", "accepted", "business written"],
            negative=["outside appetite", "not written", "excluded", "restricted"],
            neutral_default=50.0,
        ),
    }

    weighted = {k: metric_scores[k] * METRIC_WEIGHTS[k] for k in METRIC_WEIGHTS}
    risk_score = round(sum(weighted.values()), 2)

    available_signals = 0
    for signal in [client_lr, avg_loss_ratio, percentile, claims_freq_percentile, severity_percentile, incurred_percentile, web_summary, intranet_summary, doc_insights]:
        if signal not in [None, "", []]:
            available_signals += 1
    confidence = round(min(98.0, 45.0 + available_signals * 10.0), 1)

    if any("sanctions hit" in h.lower() for h in hard_rules):
        decision = "DECLINE"
    elif any("above decline threshold" in h.lower() or "hard-stop" in h.lower() for h in hard_rules):
        decision = "DECLINE"
    elif any("exceeds delegated authority" in h.lower() for h in hard_rules):
        decision = "REFER"
    elif risk_score < SCORE_BANDS["write"]:
        decision = "WRITE"
    elif risk_score <= SCORE_BANDS["conditional"]:
        decision = "WRITE_WITH_CONDITIONS"
    elif risk_score <= SCORE_BANDS["refer"]:
        decision = "REFER"
    else:
        decision = "DECLINE"

    top_risks = sorted(weighted.items(), key=lambda x: x[1], reverse=True)[:3]
    top_positives = sorted(weighted.items(), key=lambda x: x[1])[:3]
    key_risk_drivers = [f"{k.replace('_', ' ').title()} ({metric_scores[k]:.1f}/100)" for k, _ in top_risks]
    key_positive_factors = [f"{k.replace('_', ' ').title()} ({metric_scores[k]:.1f}/100)" for k, _ in top_positives]

    conditions = _extract_conditions(doc_insights, intranet_summary, web_summary, risk_score)
    if decision not in {"WRITE_WITH_CONDITIONS", "REFER"}:
        conditions = []

    decision_doc_type = {
        "WRITE": "AUTO_EOI",
        "WRITE_WITH_CONDITIONS": "CONDITIONAL_EOI",
        "REFER": "UNDERWRITING_MEMO",
        "DECLINE": "DECLINE_LETTER",
    }.get(decision, "UNDERWRITING_MEMO")

    return {
        "decision": decision,
        "decision_document_type": decision_doc_type,
        "risk_score": risk_score,
        "confidence_score": confidence,
        "web_risk_score": web_risk_score,
        "hard_rule_triggered": bool(hard_rules),
        "hard_rule_hits": hard_rules,
        "key_positive_factors": key_positive_factors,
        "key_risk_drivers": key_risk_drivers,
        "conditions": conditions,
        "metric_scores": {k: round(v, 2) for k, v in metric_scores.items()},
        "weighted_contributions": {k: round(v, 2) for k, v in weighted.items()},
        "normalization_layer": {k: round(v, 2) for k, v in metric_scores.items()},
    }


def _run_external_search(prompt: str) -> dict[str, Any]:
    try:
        from agents.serp_agent import serp_node

        out = serp_node({"user_prompt": prompt})
        return {
            "web_links": out.get("web_links") or [],
            "general_summary": out.get("general_summary") or "",
        }
    except Exception as exc:
        return {
            "web_links": [("SERP API request failed (no links available)", str(exc))],
            "general_summary": f"External search failed: {exc}",
        }


def _build_geo_risk_prompt(user_prompt: str, risk_profile: dict[str, Any] | None, broker_fields: dict[str, str] | None) -> str:
    profile = risk_profile or {}
    fields = broker_fields or {}
    geography_parts = [
        str(fields.get("country") or "").strip(),
        str(fields.get("state") or "").strip(),
        str(profile.get("sites") or "").strip(),
    ]
    geography = ", ".join([g for g in geography_parts if g]) or "reported insured locations"
    return (
        f"{user_prompt}\n"
        f"Geography focus: {geography}.\n"
        "Using web sources, assess geography-specific risk for calamities (flood/earthquake/storm/wildfire), "
        "crime/violence, and strike/labor unrest. Include location-specific signals and severity."
    )


def _run_geo_risk_search(user_prompt: str, risk_profile: dict[str, Any] | None, broker_fields: dict[str, str] | None) -> dict[str, Any]:
    geo_prompt = _build_geo_risk_prompt(user_prompt, risk_profile, broker_fields)
    out = _run_external_search(geo_prompt)
    out["general_summary"] = _normalize_complete_sentences(str(out.get("general_summary") or ""), max_sentences=6)
    normalized_links = []
    for item in out.get("web_links") or []:
        if isinstance(item, dict):
            title = str(item.get("title") or item.get("link") or "Web Source")
            summary = str(item.get("summary") or item.get("snippet") or "")
            normalized_links.append((title, _build_geo_link_summary(title, summary)))
        elif isinstance(item, (list, tuple)) and len(item) >= 2:
            title = str(item[0] or "Web Source")
            summary = str(item[1] or "")
            normalized_links.append((title, _build_geo_link_summary(title, summary)))
        else:
            title = str(item or "Web Source")
            normalized_links.append((title, _build_geo_link_summary(title, "")))
    out["web_links"] = normalized_links
    return out | {"geo_prompt": geo_prompt}


def _infer_lob_hint(text: str) -> str | None:
    text_low = (text or "").lower()

    lob_keywords = {
        "Marine": ["marine", "cargo", "vessel", "shipping", "maritime", "hull"],
        "Casualty": ["casualty", "liability", "public liability", "general liability", "third party"],
        "Property": ["property", "fire", "building", "business interruption", "industrial all risk"],
        "Motor": ["motor", "auto", "vehicle", "fleet", "commercial vehicle"],
        "Energy": ["energy", "oil", "gas", "power", "offshore", "onshore", "renewable"],
        "Aviation": ["aviation", "aircraft", "aero", "aerospace", "airline", "airport"],
        "Financial Lines": ["financial lines", "d&o", "directors and officers", "professional indemnity", "cyber"],
        "Construction": ["construction", "contractor", "builders risk", "construction all risk", "erection all risk"],
    }

    best_lob = None
    best_score = 0
    for lob, kws in lob_keywords.items():
        score = sum(1 for kw in kws if kw in text_low)
        if score > best_score:
            best_score = score
            best_lob = lob

    return best_lob if best_score > 0 else None


def _run_intranet_insights(user_prompt: str, doc_insights: str, broker_text: str) -> dict[str, Any]:
    combined = f"{user_prompt}\n{doc_insights}\n{broker_text[:4000]}"
    lob_hint = _infer_lob_hint(combined)
    if not lob_hint:
        return {
            "intranet_summary": "",
            "intranet_sources": [],
            "intranet_doc_links": [],
            "intranet_doc_count": None,
            "intranet_lob": None,
        }

    intranet_prompt = (
        f"{user_prompt}\n"
        f"Line of business: {lob_hint}.\n"
        "Provide relevant underwriting guidelines, business written, and business not written."
    )

    try:
        out = intranet_node({"user_prompt": intranet_prompt})
        return {
            "intranet_summary": out.get("intranet_summary") or "",
            "intranet_sources": out.get("intranet_sources") or [],
            "intranet_doc_links": out.get("intranet_doc_links") or [],
            "intranet_doc_count": out.get("intranet_doc_count"),
            "intranet_lob": out.get("intranet_lob") or lob_hint,
        }
    except BaseException as exc:
        # intranet_node may raise streamlit stop-style exceptions when LOB is missing.
        return {
            "intranet_summary": f"Intranet lookup failed: {exc}",
            "intranet_sources": [],
            "intranet_doc_links": [],
            "intranet_doc_count": None,
            "intranet_lob": lob_hint,
        }


def _fallback_recommendation(
    document_summary: str,
    vanna_summary: str,
    web_summary: str,
    intranet_summary: str,
) -> str:
    combined = "\n".join([document_summary or "", vanna_summary or "", web_summary or "", intranet_summary or ""])
    low = combined.lower()
    negative_markers = [
        "do not write",
        "do not underwrite",
        "not underwrite",
        "decline",
        "excluded",
        "prohibited",
        "unacceptable risk",
    ]
    positive_markers = [
        "acceptable risk",
        "within appetite",
        "can underwrite",
        "write",
        "eligible",
    ]

    neg = sum(1 for m in negative_markers if m in low)
    pos = sum(1 for m in positive_markers if m in low)

    decision = "DO NOT WRITE" if neg >= max(1, pos) else "WRITE"
    if decision == "DO NOT WRITE":
        rationale = "material risk or exclusion indicators are stronger than supportive signals."
    else:
        rationale = "risk appears acceptable subject to normal underwriting controls."

    return (
        f"{decision}: {rationale}\n"
        f"- Internal claim history (Vanna): {_to_short_blurb(vanna_summary or 'No strong internal claim signal available.')}\n"
        f"- External insight (SERP): {_to_short_blurb(web_summary or 'No strong external signal available.')}\n"
        f"- Rules & Guidelines (Intranet): {_to_short_blurb(intranet_summary or 'No clear intranet guideline signal available.')}\n"
        f"- Broker submission (Document): {_to_short_blurb(document_summary or 'No clear broker submission signal available.')}"
    )


def _to_short_blurb(text: str, max_sentences: int = 2, max_chars: int = 280) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if not cleaned:
        return ""
    sentences = re.split(r"(?<=[.!?])\s+", cleaned)
    blurb = " ".join(sentences[:max_sentences]).strip()
    if not blurb:
        blurb = cleaned
    return blurb[:max_chars].rstrip()


def _fmt_currency(value: float | None) -> str:
    if value is None:
        return "N/A"
    return f"${value:,.2f}"


def _llm_exec_summary(agent_name: str, raw_text: str, user_prompt: str) -> str:
    prompt = f"""
You are an underwriting analyst preparing an executive snapshot.

Summarize the {agent_name} output into exactly 2 to 3 complete sentences.
Requirements:
- Use professional underwriting language.
- Keep only decision-relevant points.
- No bullet points.
- No sentence fragments.

User Prompt:
{user_prompt}

Agent Output:
{raw_text[:4000]}
"""
    try:
        raw = call_llm(prompt).strip()
        normalized = _normalize_complete_sentences(raw, max_sentences=3)
    except Exception:
        normalized = _normalize_complete_sentences(raw_text, max_sentences=3)

    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", normalized) if s.strip()]
    if len(sentences) >= 2:
        return "\n".join(sentences[:3])

    base = _ensure_sentence(normalized or raw_text or f"{agent_name} signal is available for underwriting review")
    fallback = [
        f"{agent_name} output has been incorporated into the underwriting assessment.",
        "The signal should be interpreted alongside internal policy, authority, and portfolio controls.",
    ]
    return "\n".join([base] + fallback)


def _build_vanna_underwriter_summary(sql_result: pd.DataFrame | None) -> str:
    if not isinstance(sql_result, pd.DataFrame) or sql_result.empty:
        return "Internal portfolio benchmark is unavailable for this submission."
    if "Error" in sql_result.columns:
        err = str(sql_result.iloc[0].get("Error") or "SQL benchmark failed")
        return _ensure_sentence(f"Internal portfolio benchmark could not be completed: {err}")

    row = sql_result.fillna("").iloc[0].to_dict()
    account_count = _to_float(row.get("account_count"))
    lr_pctile = _to_float(row.get("loss_ratio_percentile"))
    freq_pctile = _to_float(row.get("claims_frequency_percentile"))
    sev_pctile = _to_float(row.get("severity_percentile"))
    inc_pctile = _to_float(row.get("incurred_percentile"))
    avg_lr = _to_float(row.get("avg_loss_ratio"))
    avg_claim_freq = _to_float(row.get("avg_claims_frequency"))
    avg_severity = _to_float(row.get("avg_largest_single_loss"))
    avg_premium = _to_float(row.get("avg_ultimate_premium"))
    avg_incurred = _to_float(row.get("avg_incurred_loss"))
    client_freq = _to_float(row.get("client_claims_frequency"))
    client_severity = _to_float(row.get("client_severity"))
    client_incurred = _to_float(row.get("client_incurred_per_year"))

    lines: list[str] = []
    if account_count is not None and account_count > 0:
        lines.append(_ensure_sentence(f"Portfolio benchmark considered approximately {int(account_count)} comparable records"))
    if lr_pctile is not None:
        lines.append(_ensure_sentence(f"Client loss ratio sits at the {lr_pctile:.1f} percentile versus portfolio peers, indicating {'higher' if lr_pctile >= 60 else 'lower' if lr_pctile <= 40 else 'mid-range'} relative loss performance"))
    if freq_pctile is not None or sev_pctile is not None or inc_pctile is not None:
        fp = f"{freq_pctile:.1f}" if freq_pctile is not None else "N/A"
        sp = f"{sev_pctile:.1f}" if sev_pctile is not None else "N/A"
        ip = f"{inc_pctile:.1f}" if inc_pctile is not None else "N/A"
        lines.append(_ensure_sentence(f"Loss pattern percentiles are frequency {fp}, severity {sp}, and incurred trend {ip}"))

    portfolio_terms: list[str] = []
    if avg_premium is not None:
        portfolio_terms.append(f"average ultimate premium {_fmt_currency(avg_premium)}")
    if avg_incurred is not None:
        portfolio_terms.append(f"average incurred loss {_fmt_currency(avg_incurred)}")
    if avg_claim_freq is not None:
        portfolio_terms.append(f"average claims frequency {avg_claim_freq:.2f}")
    if avg_severity is not None:
        portfolio_terms.append(f"average severity (largest single loss) {_fmt_currency(avg_severity)}")
    if portfolio_terms:
        lines.append(_ensure_sentence(f"Portfolio benchmark shows {', '.join(portfolio_terms)}"))

    client_terms: list[str] = []
    if client_incurred is not None:
        client_terms.append(f"client incurred per year {_fmt_currency(client_incurred)}")
    if client_freq is not None:
        client_terms.append(f"client frequency {client_freq:.2f}")
    if client_severity is not None:
        client_terms.append(f"client severity {_fmt_currency(client_severity)}")
    if client_terms:
        lines.append(_ensure_sentence(f"Client-side extracted loss pattern indicates {', '.join(client_terms)}"))

    peer_stats: list[str] = []
    if avg_lr is not None:
        peer_stats.append(f"average loss ratio {avg_lr:.2f}")
    if peer_stats:
        lines.append(_ensure_sentence(f"Peer benchmark shows {', '.join(peer_stats)}"))

    if not lines:
        return "Internal SQL output returned rows, but no benchmark metrics were available for underwriting interpretation."
    return " ".join(lines[:4])


def _build_executive_snapshot(
    user_prompt: str,
    doc_insights: str,
    sql_query: str | None,
    sql_result: pd.DataFrame | None,
    web_summary: str,
    web_links: list,
    intranet_summary: str,
    decision_payload: dict[str, Any] | None = None,
) -> dict[str, str]:
    if isinstance(sql_result, pd.DataFrame) and not sql_result.empty:
        sql_context = sql_result.head(6).to_markdown(index=False)
    else:
        sql_context = "No SQL rows returned."

    links_context = "\n".join(
        [f"- {item[0]} | {item[1]}" for item in (web_links or [])[:5] if isinstance(item, (list, tuple)) and len(item) >= 2]
    ) or "No web links."
    web_link_summaries = " ".join(
        [str(item[1]) for item in (web_links or [])[:5] if isinstance(item, (list, tuple)) and len(item) >= 2 and item[1]]
    )
    web_summary_short = web_summary or web_link_summaries or "No external web summary available."

    vanna_short = _build_vanna_underwriter_summary(sql_result)
    decision_payload = decision_payload or {}
    decision = str(decision_payload.get("decision") or "WRITE")
    risk_score = decision_payload.get("risk_score")
    confidence = decision_payload.get("confidence_score")
    risks = decision_payload.get("key_risk_drivers") or []
    positives = decision_payload.get("key_positive_factors") or []
    conditions = decision_payload.get("conditions") or []
    hard_rules = decision_payload.get("hard_rule_hits") or []

    final_lines = [
        f"Decision: {decision}",
        f"Risk Score: {risk_score if risk_score is not None else 'N/A'} / 100",
        f"Confidence: {confidence if confidence is not None else 'N/A'}%",
        f"Key Risk Drivers: {', '.join(risks[:3]) if risks else 'Not identified'}",
        f"Key Positive Factors: {', '.join(positives[:3]) if positives else 'Not identified'}",
    ]
    if hard_rules:
        final_lines.append(f"Hard Rule Hits: {'; '.join(hard_rules[:2])}")
    if conditions:
        final_lines.append(f"Conditions: {'; '.join(conditions[:2])}")

    fallback = {
        "document_agent_summary": _llm_exec_summary("Document Agent", doc_insights or "No document insight available.", user_prompt),
        "vanna_agent_summary": _llm_exec_summary("SQL Agent", vanna_short, user_prompt),
        "web_agent_summary": _llm_exec_summary("Web Agent", web_summary_short, user_prompt),
        "intranet_agent_summary": _llm_exec_summary("Intranet Agent", intranet_summary or "No intranet policy insight available.", user_prompt),
        "final_recommendation": "\n".join(final_lines),
        "decision": decision,
        "risk_score": str(risk_score if risk_score is not None else ""),
        "confidence_score": str(confidence if confidence is not None else ""),
    }

    return fallback


def _load_template_text(template_path: str) -> str:
    ext = Path(template_path).suffix.lower()
    if ext == ".docx":
        doc = Document(template_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    if ext == ".pdf":
        if PdfReader is None:
            raise RuntimeError("PDF reader library not installed. Install `pypdf`.")
        reader = PdfReader(template_path)
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
        if text.strip():
            return text
        raise ValueError(f"Template PDF has no extractable text: {template_path}")
    raise ValueError(f"Unsupported template format: {template_path}")


def _resolve_template_path(template_path: str | None) -> Path:
    candidates: list[Path] = []
    if template_path:
        candidates.append(Path(template_path))

    env_template = os.getenv("EOI_TEMPLATE_PATH")
    if env_template:
        candidates.append(Path(env_template))

    candidates.append(Path(DEFAULT_EOI_TEMPLATE_PATH))
    candidates.append(Path(PDF_EOI_TEMPLATE_PATH))

    resolved_candidates: list[Path] = []
    for candidate in candidates:
        p = candidate.expanduser()
        if not p.is_absolute():
            p = PROJECT_ROOT / p
        p = p.resolve()
        resolved_candidates.append(p)
        if p.exists():
            return p

    attempted = resolved_candidates[0] if resolved_candidates else Path(DEFAULT_EOI_TEMPLATE_PATH)
    raise FileNotFoundError(f"EOI template not found: {attempted}")


def _build_filled_form_text(user_prompt: str, eoi_state: dict[str, Any], template_text: str) -> str:
    sql_result = eoi_state.get("sql_result")
    if isinstance(sql_result, pd.DataFrame) and not sql_result.empty:
        sql_context = sql_result.head(12).to_markdown(index=False)
    else:
        sql_context = "No structured internal output available."

    web_links = eoi_state.get("web_links") or []
    web_context = "\n".join([f"- {link} | {summary}" for link, summary in web_links[:5]]) or "No web links."

    doc_insights = eoi_state.get("eoi_doc_insights") or "No document insights available."
    web_summary = eoi_state.get("general_summary") or "No web summary available."
    intranet_summary = eoi_state.get("intranet_summary") or "No intranet policy insights available."
    broker_fields = eoi_state.get("eoi_broker_fields") or {}

    prompt = f"""
You are filling an insurance EOI form.

Use ONLY the evidence below:
1) Broker submission insights
2) Internal SQL insights
3) External web insights
4) Intranet policy and underwriting guideline insights
5) Extracted broker fields (highest priority for matching labels)

If any field is missing, fill it with "Not Provided".
Preserve the exact form structure and section order from TEMPLATE below.

Formatting rules:
- Keep labels exactly as-is.
- Replace blank lines/underscores with values.
- For checkboxes, mark selected options with "\\u2611" and unselected with "\\u2610".
- In "Type of Entity", select exactly one option.
- Return plain text only. No markdown, no explanations.

EXTRACTED BROKER FIELDS (PRIORITY):
{json.dumps(broker_fields, indent=2)}

USER PROMPT:
{user_prompt}

BROKER SUBMISSION INSIGHTS:
{doc_insights}

INTERNAL SQL SUMMARY:
Query: {eoi_state.get('sql_query') or 'Not Provided'}
Top rows:
{sql_context}

EXTERNAL WEB SUMMARY:
{web_summary}

INTRANET POLICY INSIGHTS:
{intranet_summary}

EXTERNAL LINKS:
{web_context}

TEMPLATE:
{template_text}
"""

    try:
        filled = call_llm(prompt).strip()
    except Exception:
        filled = ""

    # Fallback: keep deterministic minimum form body if LLM fill fails/returns empty.
    if not filled:
        lines = [
            "Insurance Expression of Interest",
            "",
            "1. Applicant Information",
        ]
        for key, label in FIELD_LABELS.items():
            value = str(broker_fields.get(key) or "Not Provided")
            lines.append(f"{label}: {value}")
        lines.extend(
            [
                "",
                "2. Submission Context",
                f"Prompt: {user_prompt or 'Not Provided'}",
                f"Document Insights: {_to_short_blurb(doc_insights)}",
            ]
        )
        filled = "\n".join(lines)

    return filled


def _replace_section_block(lines: list[str], heading_patterns: list[str], new_block: list[str]) -> list[str]:
    start_idx = None
    for i, ln in enumerate(lines):
        low = ln.strip().lower()
        if any(re.fullmatch(pat, low) for pat in heading_patterns):
            start_idx = i
            break

    if start_idx is None:
        if lines and lines[-1].strip():
            lines.append("")
        return lines + new_block

    end_idx = len(lines)
    for j in range(start_idx + 1, len(lines)):
        candidate = lines[j].strip()
        low = candidate.lower()
        if re.match(r"^\d+\.\s", candidate):
            end_idx = j
            break
        if low.startswith("disclaimer"):
            end_idx = j
            break
        if low.startswith("authorized signature"):
            end_idx = j
            break

    return lines[:start_idx] + new_block + lines[end_idx:]


def _today_str() -> str:
    return datetime.now().strftime("%d %b %Y")


def _enforce_today_date(filled_text: str) -> str:
    lines = filled_text.splitlines()
    out = []
    date_set = False

    for line in lines:
        if re.match(r"^\s*date\s*:", line, flags=re.IGNORECASE):
            out.append(f"Date: {_today_str()}")
            date_set = True
        else:
            out.append(line)

    if not date_set:
        inserted = False
        for idx, ln in enumerate(out):
            if re.match(r"^\s*1\.\s", ln):
                out.insert(idx, f"Date: {_today_str()}")
                out.insert(idx, "")
                inserted = True
                break
        if not inserted:
            out.append("")
            out.append(f"Date: {_today_str()}")

    return "\n".join(out)


def _ensure_disclaimer_and_signature(filled_text: str) -> str:
    lines = filled_text.splitlines()
    low_all = "\n".join(lines).lower()

    if "disclaimer:" not in low_all:
        lines.extend(
            [
                "",
                "Disclaimer: This document is generated from submitted and referenced data. Final underwriting decision remains subject to internal approval and policy terms.",
            ]
        )
    if "authorized signature:" not in low_all:
        lines.extend(
            [
                "",
                "Authorized Signature: ____________________",
                "Name: ____________________",
                f"Date: {_today_str()}",
            ]
        )
    return "\n".join(lines)


def _append_runtime_insights(filled_text: str, user_prompt: str, eoi_state: dict[str, Any]) -> str:
    decision_payload = eoi_state.get("eoi_decision_payload") or {}
    risk_profile = eoi_state.get("eoi_risk_profile") or {}
    sql_result = eoi_state.get("sql_result")
    sql_query = eoi_state.get("sql_query") or "Not Provided"

    sql_snapshot = "Not available"
    if isinstance(sql_result, pd.DataFrame) and not sql_result.empty:
        if "Error" in sql_result.columns:
            sql_snapshot = str(sql_result.iloc[0].get("Error") or "SQL error")
        else:
            row = sql_result.fillna("").iloc[0].to_dict()
            parts = [f"{k}: {row[k]}" for k in list(row.keys())[:6] if str(row[k]).strip()]
            sql_snapshot = "; ".join(parts) if parts else "No rows"

    hard_rules = decision_payload.get("hard_rule_hits") or []
    conditions = decision_payload.get("conditions") or []
    risk_drivers = decision_payload.get("key_risk_drivers") or []
    positives = decision_payload.get("key_positive_factors") or []

    insight_lines = [
        "AI-DERIVED RISK INSIGHTS",
        f"User Prompt: {user_prompt or 'Not Provided'}",
        f"Decision: {decision_payload.get('decision', eoi_state.get('eoi_decision', 'Not Provided'))}",
        f"Risk Score: {decision_payload.get('risk_score', eoi_state.get('eoi_risk_score', 'Not Provided'))}",
        f"Confidence Score: {decision_payload.get('confidence_score', eoi_state.get('eoi_confidence_score', 'Not Provided'))}",
        "",
        "Risk Profile JSON",
        f"LOB: {risk_profile.get('lob', 'Not specified')}",
        f"TIV: {risk_profile.get('tiv', 'Not specified')}",
        f"Turnover: {risk_profile.get('turnover', 'Not specified')}",
        f"Sites: {risk_profile.get('sites', 'Not specified')}",
        "",
        "Agent Evidence Summary",
        f"Document Agent: {_to_short_blurb(eoi_state.get('eoi_doc_insights') or 'Not available.')}",
        f"SQL Query: {sql_query}",
        f"SQL Snapshot: {_to_short_blurb(sql_snapshot, max_sentences=3, max_chars=500)}",
        f"Web Agent: {_to_short_blurb(eoi_state.get('general_summary') or 'Not available.', max_sentences=3, max_chars=500)}",
        f"Intranet Agent: {_to_short_blurb(eoi_state.get('intranet_summary') or 'Not available.', max_sentences=3, max_chars=500)}",
        "",
        "Decision Drivers",
        f"Key Positive Factors: {', '.join(positives) if positives else 'Not identified'}",
        f"Key Risk Drivers: {', '.join(risk_drivers) if risk_drivers else 'Not identified'}",
        f"Hard Rule Hits: {'; '.join(hard_rules) if hard_rules else 'None'}",
        f"Conditions: {'; '.join(conditions) if conditions else 'None'}",
    ]
    lines = filled_text.splitlines()
    replaced = _replace_section_block(
        lines,
        heading_patterns=[
            r"ai-derived risk insights",
            r"\d+\.\s*ai-derived risk insights",
            r"ai underwriting insights",
            r"\d+\.\s*ai underwriting insights",
        ],
        new_block=insight_lines,
    )
    return "\n".join(replaced)


def _override_field_line(line: str, label: str, value: str) -> str:
    if not value:
        return line
    pattern = rf"^\s*{re.escape(label)}\s*:"
    if re.match(pattern, line, flags=re.IGNORECASE):
        return f"{label}: {value}"
    return line


def _enforce_extracted_fields(filled_text: str, broker_fields: dict[str, str]) -> str:
    lines = filled_text.splitlines()
    for i, line in enumerate(lines):
        updated = line
        for key, label in FIELD_LABELS.items():
            updated = _override_field_line(updated, label, broker_fields.get(key, ""))
        lines[i] = updated
    return "\n".join(lines)


def _format_vanna_claims_history(eoi_state: dict[str, Any]) -> str:
    sql_result = eoi_state.get("sql_result")
    if not isinstance(sql_result, pd.DataFrame) or sql_result.empty:
        return ""
    if "Error" in sql_result.columns:
        return ""

    # Build compact row summaries from Vanna output for claim history section.
    top_rows = sql_result.head(3).fillna("")
    row_summaries: list[str] = []
    for _, row in top_rows.iterrows():
        pairs = []
        for col in top_rows.columns[:4]:
            val = str(row[col]).strip()
            if not val:
                continue
            pairs.append(f"{col}: {val}")
        if pairs:
            row_summaries.append(", ".join(pairs))

    return " | ".join(row_summaries)


def _enforce_declaration_and_claims(filled_text: str) -> str:
    lines = filled_text.splitlines()
    out: list[str] = []
    declaration_seen = False

    i = 0
    while i < len(lines):
        line = lines[i]
        low = line.strip().lower()

        if low.startswith("i/we hereby declare that"):
            out.append(DECLARATION_TEXT)
            declaration_seen = True
            i += 1
            continue

        if low.startswith("claims history (last 3") and i + 1 < len(lines):
            next_line = lines[i + 1].strip().lower()
            out.append(line)
            if next_line == "not provided":
                i += 2
                continue
            i += 1
            continue

        out.append(line)
        i += 1

    if not declaration_seen:
        for idx, line in enumerate(out):
            if line.strip().lower().startswith("6. declaration"):
                out.insert(idx + 1, "")
                out.insert(idx + 2, DECLARATION_TEXT)
                break

    return "\n".join(out)


def _infer_entity_type(user_prompt: str, eoi_state: dict[str, Any]) -> str:
    text = f"{user_prompt}\n{eoi_state.get('eoi_doc_insights','')}".lower()

    if any(k in text for k in ["broker", "broking", "brokerage", "intermediary"]):
        return "Broker"
    if any(k in text for k in ["vendor", "partner", "supplier"]):
        return "Vendor / Partner"
    if any(k in text for k in ["corporate", "company", "corporation", "inc", "llc", "ltd", "private limited"]):
        return "Corporate"
    if any(k in text for k in ["individual", "person", "sole proprietor", "sole-proprietor"]):
        return "Individual"
    return "Other"


def _set_checkbox_mark(line: str, checked: bool) -> str:
    mark = CHECKED_BOX if checked else UNCHECKED_BOX
    normalized = line.replace("\\u2611", CHECKED_BOX).replace("\\u2610", UNCHECKED_BOX)
    cleaned = re.sub(rf"[{CHECKED_BOX}{UNCHECKED_BOX}]", "", normalized).strip()
    return f"{mark} {cleaned}" if cleaned else mark


def _enforce_type_entity_checkboxes(filled_text: str, selected_entity: str) -> str:
    lines = filled_text.splitlines()
    in_entity_block = False

    for i, line in enumerate(lines):
        lower = line.lower().strip()

        if "type of entity" in lower:
            in_entity_block = True
            continue

        if in_entity_block and lower.startswith("registration / license number"):
            in_entity_block = False

        if not in_entity_block:
            continue

        if "individual" in lower:
            lines[i] = _set_checkbox_mark(line, selected_entity == "Individual")
        elif "corporate" in lower:
            lines[i] = _set_checkbox_mark(line, selected_entity == "Corporate")
        elif "broker" in lower:
            lines[i] = _set_checkbox_mark(line, selected_entity == "Broker")
        elif "vendor" in lower and "partner" in lower:
            lines[i] = _set_checkbox_mark(line, selected_entity == "Vendor / Partner")
        elif "other" in lower:
            lines[i] = _set_checkbox_mark(line, selected_entity == "Other")

    return "\n".join(lines)


def _normalize_checkbox_lines(filled_text: str) -> str:
    out = []
    for raw in filled_text.splitlines():
        line = raw.replace("\\u2611", CHECKED_BOX).replace("\\u2610", UNCHECKED_BOX).strip()

        if CHECKED_BOX in line or UNCHECKED_BOX in line:
            checked = CHECKED_BOX in line
            cleaned = re.sub(rf"[{CHECKED_BOX}{UNCHECKED_BOX}]", "", line).strip()
            cleaned = re.sub(r"\s+", " ", cleaned)
            line = f"{CHECKED_BOX if checked else UNCHECKED_BOX} {cleaned}" if cleaned else (CHECKED_BOX if checked else UNCHECKED_BOX)

        out.append(line)
    return "\n".join(out)


def _apply_doc_theme(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def _add_line_with_style(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.lower().startswith("insurance expression of interest"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stripped)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(13, 52, 103)
        return

    if re.match(r"^\d+\.\s", stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(25, 66, 124)
        return

    if stripped.startswith(CHECKED_BOX) or stripped.startswith(UNCHECKED_BOX):
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        if stripped.startswith(CHECKED_BOX):
            run.bold = True
            run.font.color.rgb = RGBColor(15, 118, 110)
        else:
            run.font.color.rgb = RGBColor(100, 116, 139)
        return

    if ":" in stripped:
        label, value = stripped.split(":", 1)
        p = doc.add_paragraph()
        k = p.add_run(f"{label.strip()}: ")
        k.bold = True
        k.font.color.rgb = RGBColor(31, 41, 55)

        value_text = value.strip() or "Not Provided"
        v = p.add_run(value_text)
        v.font.color.rgb = RGBColor(16, 87, 156)
        return

    p = doc.add_paragraph()
    run = p.add_run(stripped)
    run.font.color.rgb = RGBColor(55, 65, 81)


def _claims_history_df_from_state(eoi_state: dict[str, Any] | None) -> pd.DataFrame | None:
    if not eoi_state:
        return None
    sql_result = eoi_state.get("sql_result")
    if not isinstance(sql_result, pd.DataFrame) or sql_result.empty:
        return None
    if "Error" in sql_result.columns:
        return None
    drop_cols = [c for c in sql_result.columns if str(c).strip().lower() in SQL_HIDDEN_OUTPUT_COLUMNS]
    view = sql_result.drop(columns=drop_cols, errors="ignore")
    return view.head(8).copy()


def _add_claims_history_table(doc: Document, df: pd.DataFrame) -> None:
    if df.empty:
        return

    view = df.iloc[:, : min(len(df.columns), 6)].fillna("")
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(view.columns):
        hdr_cells[i].text = str(col)

    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row.tolist()):
            cells[i].text = str(val)


def _build_styled_eoi_doc(filled_text: str, eoi_state: dict[str, Any] | None = None) -> bytes:
    doc = Document()
    _apply_doc_theme(doc)
    claims_df = _claims_history_df_from_state(eoi_state)

    for line in filled_text.splitlines():
        if line.strip().lower().startswith("claims history (last 3"):
            label = "Claims History (last 3–5 years)"
            value = "See Vanna claims table below." if claims_df is not None else line.split(":", 1)[-1].strip()
            _add_line_with_style(doc, f"{label}: {value}")
            if claims_df is not None:
                _add_claims_history_table(doc, claims_df)
            continue
        _add_line_with_style(doc, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _enforce_calibri_font(doc: Document) -> None:
    for p in _iter_all_paragraphs(doc):
        for r in p.runs:
            r.font.name = "Calibri"


def _set_paragraph_text_preserve_format(p, text: str, force_bold: bool | None = None) -> None:
    if not p.runs:
        run = p.add_run("")
        if force_bold is not None:
            run.bold = force_bold
    p.runs[0].text = text
    if force_bold is not None:
        p.runs[0].bold = force_bold
    # Keep non-heading/body lines visually consistent even if template paragraph style varies.
    if force_bold is False:
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11)
    for r in p.runs[1:]:
        r.text = ""


def _replace_prefixed_line(doc: Document, prefix: str, value: str) -> bool:
    low_prefix = prefix.lower()
    for p in _iter_all_paragraphs(doc):
        txt = (p.text or "").strip()
        if txt.lower().startswith(low_prefix):
            _set_paragraph_text_preserve_format(p, f"{prefix}{value}")
            return True
    return False


def _replace_heading_block(doc: Document, heading: str, lines: list[str]) -> None:
    def _norm(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

    def _heading_text(base_heading: str) -> str:
        clean = re.sub(r"^\d+\.\s*", "", base_heading.strip())
        mapping = {
            "proposaloverview": "1. PROPOSAL OVERVIEW",
            "indicativeterms": "2. INDICATIVE TERMS",
            "aiderivedriskinsights": "3. AI-DERIVED RISK INSIGHTS",
            "specificconditionssubjectivities": "4. SPECIFIC CONDITIONS & SUBJECTIVITIES",
            "nextsteps": "5. NEXT STEPS",
            "disclaimer": "6. Disclaimer",
            "authorizedsignature": "7. Authorized Signature",
        }
        return mapping.get(_norm(clean), clean)

    def _set_heading_para(p, text: str) -> None:
        _set_paragraph_text_preserve_format(p, text, force_bold=True)

    paras = list(_iter_all_paragraphs(doc))
    start = None
    target_norm = _norm(heading)
    for i, p in enumerate(paras):
        candidate_norm = _norm((p.text or "").strip())
        if candidate_norm and (target_norm in candidate_norm or candidate_norm in target_norm):
            start = i
            break
    if start is None:
        anchor = None
        for p in paras:
            n = _norm((p.text or "").strip())
            if n in {"disclaimer", "authorizedsignature"}:
                anchor = p
                break
        if anchor is None:
            doc.add_paragraph("")
            h = doc.add_paragraph("")
            _set_heading_para(h, _heading_text(heading))
            for line in lines:
                np = doc.add_paragraph("")
                _set_paragraph_text_preserve_format(np, line, force_bold=False)
        else:
            h = anchor.insert_paragraph_before("")
            h.style = anchor.style
            _set_heading_para(h, _heading_text(heading))
            for line in reversed(lines):
                np = anchor.insert_paragraph_before("")
                np.style = anchor.style
                _set_paragraph_text_preserve_format(np, line, force_bold=False)
        return

    end = len(paras)
    for j in range(start + 1, len(paras)):
        t_norm = _norm((paras[j].text or "").strip())
        if t_norm in {
            "proposaloverview",
            "indicativeterms",
            "aiderivedriskinsights",
            "specificconditionssubjectivities",
            "nextsteps",
            "disclaimer",
            "authorizedsignature",
        }:
            end = j
            break

    _set_heading_para(paras[start], _heading_text(heading))
    write_index = start + 1
    for line in lines:
        if write_index < end:
            _set_paragraph_text_preserve_format(paras[write_index], line, force_bold=False)
            write_index += 1
        else:
            np = paras[end - 1].insert_paragraph_before("")
            np.style = paras[end - 1].style
            _set_paragraph_text_preserve_format(np, line, force_bold=False)

    while write_index < end:
        _set_paragraph_text_preserve_format(paras[write_index], "", force_bold=False)
        write_index += 1


def _sql_dataframe_to_table_lines(df: pd.DataFrame, max_rows: int = 5) -> list[str]:
    if not isinstance(df, pd.DataFrame) or df.empty or "Error" in df.columns:
        return ["| Status |", "| --- |", "| No structured SQL output available. |"]

    drop_cols = [c for c in df.columns if str(c).strip().lower() in SQL_HIDDEN_OUTPUT_COLUMNS]
    view = df.drop(columns=drop_cols, errors="ignore").head(max_rows).fillna("").copy()
    cols = [str(c) for c in view.columns]
    lines = []
    lines.append("| " + " | ".join(cols) + " |")
    lines.append("| " + " | ".join(["---"] * len(cols)) + " |")
    for _, row in view.iterrows():
        vals = [str(row[c]).replace("\n", " ").strip() for c in view.columns]
        lines.append("| " + " | ".join(vals) + " |")
    return lines


def _sql_dataframe_for_doc_table(df: pd.DataFrame | None, max_rows: int = 5) -> tuple[list[str], list[list[str]]]:
    if not isinstance(df, pd.DataFrame) or df.empty or "Error" in df.columns:
        return ["Status"], [["No structured SQL output available."]]

    def _pretty_sql_col(col: str) -> str:
        overrides = {
            "account_count": "Account Count",
            "avg_loss_ratio": "Avg Loss Ratio",
            "avg_claims_frequency": "Avg Claims Frequency",
            "avg_largest_single_loss": "Avg Largest Single Loss",
            "avg_ultimate_premium": "Avg Ultimate Premium",
            "avg_incurred_loss": "Avg Incurred Loss",
            "loss_ratio_percentile": "Loss Ratio Percentile",
            "claims_frequency_percentile": "Claims Frequency Percentile",
            "severity_percentile": "Severity Percentile",
            "incurred_percentile": "Incurred Percentile",
            "client_loss_ratio": "Client Loss Ratio",
            "client_claims_frequency": "Client Claims Frequency",
            "client_severity": "Client Severity",
            "client_incurred_per_year": "Client Incurred Per Year",
        }
        key = str(col).strip().lower()
        return overrides.get(key, str(col).replace("_", " ").title())

    def _format_sql_cell(col: str, value: Any) -> str:
        if value is None or str(value).strip() == "":
            return ""
        txt = str(value).strip()
        num = _to_float(value)
        if num is None:
            return txt

        col_low = str(col).lower()
        money_keys = ("premium", "incurred", "loss", "limit", "tiv", "turnover")
        if "percentile" in col_low:
            return f"{num:.2f}"
        if "ratio" in col_low:
            pct = num * 100.0 if 0 <= num <= 1 else num
            return f"{pct:.2f}%"
        if any(k in col_low for k in money_keys):
            return f"{num:,.2f}"
        if abs(num - round(num)) < 1e-9:
            return f"{int(round(num))}"
        return f"{num:.2f}"

    keep_cols = [c for c in df.columns if str(c).strip().lower() not in SQL_HIDDEN_OUTPUT_COLUMNS]
    view = df[keep_cols].head(max_rows).fillna("").copy() if keep_cols else pd.DataFrame()
    if view.empty:
        return ["Status"], [["No structured SQL output available."]]
    headers = [_pretty_sql_col(str(c)) for c in view.columns]
    rows: list[list[str]] = []
    for _, row in view.iterrows():
        rows.append([_format_sql_cell(str(c), row[c]) for c in view.columns])
    return headers, rows


def _insert_table_after_paragraph(doc: Document, paragraph, headers: list[str], rows: list[list[str]]) -> None:
    col_count = max(1, len(headers))
    table = doc.add_table(rows=1 + max(1, len(rows)), cols=col_count)
    # Some templates don't include "Table Grid"; fall back gracefully.
    style_candidates = ["Table Grid", "Light Grid", "Normal Table"]
    for style_name in style_candidates:
        try:
            table.style = style_name
            break
        except Exception:
            continue

    for idx, h in enumerate(headers):
        table.cell(0, idx).text = str(h)
        for p in table.cell(0, idx).paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10)

    if rows:
        for r_idx, row in enumerate(rows, start=1):
            for c_idx in range(col_count):
                val = row[c_idx] if c_idx < len(row) else ""
                table.cell(r_idx, c_idx).text = str(val)
                for p in table.cell(r_idx, c_idx).paragraphs:
                    for r in p.runs:
                        r.bold = False
                        r.font.name = "Calibri"
                        r.font.size = Pt(10)
    else:
        table.cell(1, 0).text = "No structured SQL output available."

    # Apply full borders to all cells.
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            for edge in ("top", "left", "bottom", "right"):
                tag = qn(f"w:{edge}")
                element = tcBorders.find(tag)
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    tcBorders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "8")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "000000")

    paragraph._p.addnext(table._tbl)


def _ensure_disclaimer_heading_spacing(doc: Document) -> None:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    def _insert_paragraph_after(paragraph, text: str, bold: bool = False):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        para = Paragraph(new_p, paragraph._parent)
        _set_paragraph_text_preserve_format(para, text, force_bold=bold)
        return para

    paras = list(_iter_all_paragraphs(doc))
    for p in paras:
        txt = (p.text or "").strip()
        low = txt.lower()
        if low.startswith("6. disclaimer"):
            prev_blank = p.insert_paragraph_before("")
            _set_paragraph_text_preserve_format(prev_blank, "", force_bold=False)
            if not p.runs:
                p.add_run(p.text)
            for r in p.runs:
                r.bold = True
            return
        if low.startswith("disclaimer:"):
            body = txt.split(":", 1)[1].strip() if ":" in txt else ""
            space_para = p.insert_paragraph_before("")
            _set_paragraph_text_preserve_format(space_para, "", force_bold=False)
            _set_paragraph_text_preserve_format(p, "6. Disclaimer", force_bold=True)
            _insert_paragraph_after(
                p,
                body or "This document is an Expression of Interest only and does not constitute a contract of insurance.",
                bold=False,
            )
            return

    # Missing disclaimer: append with spacing and bold heading.
    doc.add_paragraph("")
    h = doc.add_paragraph("")
    _set_paragraph_text_preserve_format(h, "6. Disclaimer", force_bold=True)
    doc.add_paragraph("This document is an Expression of Interest only and does not constitute a contract of insurance.")


def _build_template_based_eoi_doc(template_path: Path, user_prompt: str, eoi_state: dict[str, Any], decision_payload: dict[str, Any]) -> bytes:
    doc = Document(str(template_path))

    broker_fields = dict(eoi_state.get("eoi_broker_fields") or {})
    risk_profile = eoi_state.get("eoi_risk_profile") or {}
    web_risk = eoi_state.get("eoi_web_risk") or {}
    geo_web_summary = eoi_state.get("eoi_geo_web_summary") or ""
    conditions = decision_payload.get("conditions") or []
    hard_rules = decision_payload.get("hard_rule_hits") or []
    sql_result = eoi_state.get("sql_result")

    insured_name = broker_fields.get("registered_address") or "Prospective Insured"
    lob = risk_profile.get("lob") or "Not specified"
    proposed_limit = broker_fields.get("expected_sum_insured") or "Not Provided"
    proposed_retention = "To be finalized as per underwriting review."
    territory = risk_profile.get("sites") or "As per submission"

    _replace_prefixed_line(doc, "Date: ", _today_str())
    _replace_prefixed_line(doc, "Subject: ", f"{user_prompt or 'Insurance Program Submission'}")
    _replace_prefixed_line(doc, "Insured: ", str(insured_name))
    _replace_prefixed_line(doc, "Line of Business: ", str(lob))
    _replace_prefixed_line(doc, "Proposed Limit: ", str(proposed_limit))
    _replace_prefixed_line(doc, "Proposed Retention: ", proposed_retention)
    _replace_prefixed_line(doc, "Territory: ", str(territory))

    sql_headers, sql_rows = _sql_dataframe_for_doc_table(sql_result, max_rows=8)

    doc_agent_full = _llm_exec_summary(
        "Document Agent",
        str(eoi_state.get("eoi_doc_insights") or "Not available.").strip(),
        user_prompt,
    ).replace("\n", " ")
    web_agent_short = _llm_exec_summary(
        "Web Agent",
        str(eoi_state.get("general_summary") or "Not available.").strip(),
        user_prompt,
    ).replace("\n", " ")
    intranet_agent_short = _llm_exec_summary(
        "Intranet Agent",
        str(eoi_state.get("intranet_summary") or "Not available.").strip(),
        user_prompt,
    ).replace("\n", " ")

    geo_tokens = web_risk.get("geo_tokens") or []
    detected = web_risk.get("detected_hazards") or {}
    detected_labels = [k.title() for k, v in detected.items() if v]
    if not detected_labels:
        detected_labels = ["None flagged"]

    ai_lines = [
        f"Our underwriting engine has assigned Risk Score {decision_payload.get('risk_score', 'N/A')} with decision {decision_payload.get('decision', 'N/A')}.",
        f"Confidence Score: {decision_payload.get('confidence_score', 'N/A')}.",
        f"Document Agent Summary: {doc_agent_full}",
        f"Geography Risk Score (Web Agent): {web_risk.get('score', 'N/A')} ({web_risk.get('level', 'N/A')}).",
        f"Geographies Analyzed: {', '.join(geo_tokens) if geo_tokens else 'Not specified'}.",
        f"Geography Hazard Flags: {', '.join(detected_labels)}.",
        f"Geo Risk SERP Summary: {_to_short_blurb(geo_web_summary or 'Not available.', max_sentences=3, max_chars=500)}",
        "SQL Agent Output (Top Rows):",
    ]
    ai_lines.extend([
        f"Web Agent: {web_agent_short}",
        f"Intranet Agent: {intranet_agent_short}",
    ])
    if hard_rules:
        ai_lines.append(f"Hard Rule Hits: {'; '.join(hard_rules)}.")
    _replace_heading_block(doc, "AI-DERIVED RISK INSIGHTS", ai_lines)

    # Strictly render SQL output as a Word table (header row + data rows)
    sql_anchor = None
    for p in _iter_all_paragraphs(doc):
        if (p.text or "").strip().lower().startswith("sql agent output (top rows):"):
            sql_anchor = p
            break
    if sql_anchor is not None:
        _insert_table_after_paragraph(doc, sql_anchor, sql_headers, sql_rows)

    conditions_block = conditions[:5] if conditions else ["No additional subjectivities triggered by hard rules."]
    _replace_heading_block(doc, "SPECIFIC CONDITIONS & SUBJECTIVITIES", conditions_block)

    text_blob = "\n".join((p.text or "") for p in _iter_all_paragraphs(doc)).lower()
    _ensure_disclaimer_heading_spacing(doc)
    if "authorized signature:" not in text_blob:
        doc.add_paragraph("Authorized Signature: ____________________")
        doc.add_paragraph("Name: ____________________")
        doc.add_paragraph(f"Date: {_today_str()}")
    _enforce_calibri_font(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _append_conditions_section(filled_text: str, conditions: list[str]) -> str:
    if not conditions:
        return filled_text
    section = ["", "7. Underwriting Conditions"]
    for idx, cond in enumerate(conditions, 1):
        section.append(f"{idx}. {cond}")
    return filled_text + "\n" + "\n".join(section)


def _build_underwriting_memo_doc(user_prompt: str, eoi_state: dict[str, Any], decision_payload: dict[str, Any]) -> bytes:
    doc = Document()
    _apply_doc_theme(doc)
    _add_line_with_style(doc, "Underwriting Memo - Human Review Required")
    _add_line_with_style(doc, f"Decision: {decision_payload.get('decision', 'REFER')}")
    _add_line_with_style(doc, f"Risk Score: {decision_payload.get('risk_score', 'N/A')}")
    _add_line_with_style(doc, f"Confidence Score: {decision_payload.get('confidence_score', 'N/A')}")
    _add_line_with_style(doc, f"User Prompt: {user_prompt}")
    _add_line_with_style(doc, "")
    _add_line_with_style(doc, "Key Risk Drivers:")
    for item in decision_payload.get("key_risk_drivers", []):
        _add_line_with_style(doc, f"- {item}")
    _add_line_with_style(doc, "")
    _add_line_with_style(doc, "Conditions / Referral Notes:")
    for item in decision_payload.get("conditions", []) or ["Senior underwriting authority review required before bind."]:
        _add_line_with_style(doc, f"- {item}")
    _add_line_with_style(doc, "")
    _add_line_with_style(doc, "Source Summary:")
    _add_line_with_style(doc, f"Document: {eoi_state.get('eoi_doc_insights', 'Not available')}")
    _add_line_with_style(doc, f"Web: {eoi_state.get('general_summary', 'Not available')}")
    _add_line_with_style(doc, f"Intranet: {eoi_state.get('intranet_summary', 'Not available')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _build_decline_letter_doc(eoi_state: dict[str, Any], decision_payload: dict[str, Any]) -> bytes:
    doc = Document()
    _apply_doc_theme(doc)
    _add_line_with_style(doc, "Decline Letter")
    _add_line_with_style(doc, "")
    _add_line_with_style(
        doc,
        "Thank you for your submission. After review of internal guidelines, portfolio indicators, and external risk intelligence, we are unable to offer terms at this time.",
    )
    _add_line_with_style(doc, "")
    _add_line_with_style(doc, f"Decision: {decision_payload.get('decision', 'DECLINE')}")
    _add_line_with_style(doc, f"Risk Score: {decision_payload.get('risk_score', 'N/A')}")
    _add_line_with_style(doc, f"Confidence Score: {decision_payload.get('confidence_score', 'N/A')}")
    _add_line_with_style(doc, "Primary Reasons:")
    reasons = decision_payload.get("hard_rule_hits") or decision_payload.get("key_risk_drivers") or ["Risk profile outside underwriting appetite."]
    for r in reasons[:4]:
        _add_line_with_style(doc, f"- {r}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_eoi_document(user_prompt: str, eoi_state: dict[str, Any], template_path: str | None = None) -> tuple[bytes, str]:
    decision_payload = eoi_state.get("eoi_decision_payload") or {}
    decision = str(decision_payload.get("decision") or "WRITE")

    if decision == "DECLINE":
        return _build_decline_letter_doc(eoi_state, decision_payload), "Decline_Letter.docx"
    if decision == "REFER":
        return _build_underwriting_memo_doc(user_prompt, eoi_state, decision_payload), "Underwriting_Memo_Refer.docx"

    path = _resolve_template_path(template_path)

    if path.suffix.lower() == ".docx":
        file_name = "Generated_Insurance_EOI.docx"
        if decision == "WRITE_WITH_CONDITIONS":
            file_name = "Conditional_Insurance_EOI.docx"
        doc_bytes = _build_template_based_eoi_doc(path, user_prompt, eoi_state, decision_payload)
        return doc_bytes, file_name

    template_text = _load_template_text(str(path))
    filled_text = _build_filled_form_text(user_prompt, eoi_state, template_text)

    broker_fields = dict(eoi_state.get("eoi_broker_fields") or {})
    vanna_claims = _format_vanna_claims_history(eoi_state)
    if vanna_claims:
        broker_fields["claims_history"] = vanna_claims

    filled_text = _enforce_extracted_fields(filled_text, broker_fields)
    filled_text = _enforce_declaration_and_claims(filled_text)
    filled_text = _enforce_today_date(filled_text)
    filled_text = _append_runtime_insights(filled_text, user_prompt, eoi_state)
    filled_text = _ensure_disclaimer_and_signature(filled_text)

    selected_entity = _infer_entity_type(user_prompt, eoi_state)
    filled_text = _enforce_type_entity_checkboxes(filled_text, selected_entity)
    filled_text = _normalize_checkbox_lines(filled_text)
    if decision == "WRITE_WITH_CONDITIONS":
        filled_text = _append_conditions_section(filled_text, decision_payload.get("conditions") or [])

    doc_bytes = _build_styled_eoi_doc(filled_text, eoi_state=eoi_state)

    file_name = "Generated_Insurance_EOI.docx"
    if decision == "WRITE_WITH_CONDITIONS":
        file_name = "Conditional_Insurance_EOI.docx"
    return doc_bytes, file_name


def _empty_eoi_state(user_prompt: str = "", message: str = "EOI analysis unavailable.") -> GraphState:
    return {
        "route": "eoi",
        "vanna_prompt": build_vanna_prompt(user_prompt or ""),
        "serp_prompt": build_serp_prompt(user_prompt or ""),
        "eoi_risk_profile": {"lob": "Not specified", "tiv": "Not specified", "turnover": "Not specified", "sites": "Not specified"},
        "eoi_doc_insights": message,
        "eoi_web_risk": {"score": 50.0, "level": "MODERATE", "drivers": ["Web risk analysis unavailable."]},
        "eoi_geo_web_summary": "",
        "eoi_geo_web_links": [],
        "eoi_geo_web_prompt": "",
        "eoi_broker_fields": {},
        "sql_result": pd.DataFrame([{"Error": message}]),
        "sql_query": None,
        "web_links": [],
        "general_summary": "",
        "intranet_summary": "",
        "intranet_sources": [],
        "intranet_doc_links": [],
        "intranet_doc_count": 0,
        "intranet_lob": None,
        "eoi_decision": "REFER",
        "eoi_risk_score": 70.0,
        "eoi_confidence_score": 35.0,
        "eoi_hard_rule_hits": [],
        "eoi_conditions": ["Manual underwriter review required due to incomplete analysis signal."],
        "eoi_metric_scores": {},
        "eoi_weighted_contributions": {},
        "eoi_decision_payload": {
            "decision": "REFER",
            "decision_document_type": "UNDERWRITING_MEMO",
            "risk_score": 70.0,
            "confidence_score": 35.0,
            "web_risk_score": 50.0,
            "hard_rule_triggered": False,
            "hard_rule_hits": [],
            "key_positive_factors": [],
            "key_risk_drivers": ["Insufficient machine-readable input for full scoring."],
            "conditions": ["Manual underwriter review required due to incomplete analysis signal."],
            "metric_scores": {},
            "weighted_contributions": {},
            "normalization_layer": {},
        },
        "eoi_executive_snapshot": {
            "document_agent_summary": message,
            "vanna_agent_summary": "SQL benchmark unavailable.",
            "web_agent_summary": "Web benchmark unavailable.",
            "intranet_agent_summary": "Intranet benchmark unavailable.",
            "final_recommendation": "Decision: REFER\nRisk Score: 70 / 100\nConfidence: 35%",
            "decision": "REFER",
            "risk_score": "70",
            "confidence_score": "35",
        },
    }


def EOI_node(state: GraphState) -> GraphState:
    user_prompt = (state.get("user_prompt") or "").strip()
    if not user_prompt:
        return _empty_eoi_state(user_prompt, "Please enter a prompt.")

    try:
        uploaded_file_path = state.get("uploaded_file1_path") or state.get("uploaded_file_path")

        broker_text = ""
        if uploaded_file_path and os.path.exists(uploaded_file_path):
            try:
                broker_text = _extract_doc_text(uploaded_file_path)
            except Exception as exc:
                broker_text = ""
                state["eoi_doc_parse_error"] = str(exc)

        eoi_broker_fields = extract_broker_fields(broker_text) if broker_text else {}
        risk_profile = _extract_risk_profile_json(user_prompt, broker_text, eoi_broker_fields)

        doc_insights = "No document uploaded."
        try:
            if uploaded_file_path and os.path.exists(uploaded_file_path):
                ext = Path(uploaded_file_path).suffix.lower()
                doc_out = document_node(
                    {
                        "user_prompt": user_prompt,
                        "uploaded_file1_path": uploaded_file_path,
                        "uploaded_file1_is_excel": ext in {".xlsx", ".xls", ".csv"},
                        "uploaded_file1_is_docx": ext == ".docx",
                    }
                )
                doc_insights = doc_out.get("general_summary") or summarize_doc_with_instruction(uploaded_file_path, user_prompt)
        except Exception:
            doc_insights = summarize_doc_with_instruction(uploaded_file_path, user_prompt)

        vanna_prompt = build_vanna_prompt(user_prompt)
        serp_prompt = build_serp_prompt(user_prompt)
        sql_query, sql_result = _run_internal_sql(
            vanna_prompt,
            user_prompt=user_prompt,
            broker_fields=eoi_broker_fields,
            broker_text=broker_text,
        )
        search_result = _run_external_search(serp_prompt)
        geo_search_result = _run_geo_risk_search(
            user_prompt=user_prompt,
            risk_profile=risk_profile,
            broker_fields=eoi_broker_fields,
        )
        intranet_result = _run_intranet_insights(user_prompt, doc_insights, broker_text)
        web_risk = _compute_web_geo_risk(
            user_prompt=user_prompt,
            web_summary=geo_search_result["general_summary"],
            web_links=geo_search_result.get("web_links"),
            risk_profile=risk_profile,
            broker_fields=eoi_broker_fields,
        )
        decision_payload = _compute_risk_decision(
            user_prompt=user_prompt,
            broker_text=broker_text,
            broker_fields=eoi_broker_fields,
            risk_profile=risk_profile,
            sql_result=sql_result,
            web_summary=search_result["general_summary"],
            web_risk=web_risk,
            intranet_summary=intranet_result["intranet_summary"],
            doc_insights=doc_insights,
        )
        executive_snapshot = _build_executive_snapshot(
            user_prompt=user_prompt,
            doc_insights=doc_insights,
            sql_query=sql_query,
            sql_result=sql_result,
            web_summary=search_result["general_summary"],
            web_links=search_result["web_links"],
            intranet_summary=intranet_result["intranet_summary"],
            decision_payload=decision_payload,
        )

        return {
            "route": "eoi",
            "vanna_prompt": vanna_prompt,
            "serp_prompt": serp_prompt,
            "eoi_risk_profile": risk_profile,
            "eoi_web_risk": web_risk,
            "eoi_geo_web_summary": geo_search_result.get("general_summary"),
            "eoi_geo_web_links": geo_search_result.get("web_links"),
            "eoi_geo_web_prompt": geo_search_result.get("geo_prompt"),
            "eoi_doc_insights": doc_insights,
            "eoi_broker_fields": eoi_broker_fields,
            "sql_result": sql_result,
            "sql_query": sql_query,
            "web_links": search_result["web_links"],
            "general_summary": search_result["general_summary"],
            "intranet_summary": intranet_result["intranet_summary"],
            "intranet_sources": intranet_result["intranet_sources"],
            "intranet_doc_links": intranet_result["intranet_doc_links"],
            "intranet_doc_count": intranet_result["intranet_doc_count"],
            "intranet_lob": intranet_result["intranet_lob"],
            "eoi_decision": decision_payload.get("decision"),
            "eoi_risk_score": decision_payload.get("risk_score"),
            "eoi_confidence_score": decision_payload.get("confidence_score"),
            "eoi_hard_rule_hits": decision_payload.get("hard_rule_hits"),
            "eoi_conditions": decision_payload.get("conditions"),
            "eoi_metric_scores": decision_payload.get("metric_scores"),
            "eoi_weighted_contributions": decision_payload.get("weighted_contributions"),
            "eoi_normalization_layer": decision_payload.get("normalization_layer"),
            "eoi_hard_rule_triggered": decision_payload.get("hard_rule_triggered"),
            "eoi_decision_document_type": decision_payload.get("decision_document_type"),
            "eoi_decision_payload": decision_payload,
            "eoi_executive_snapshot": executive_snapshot,
        }
    except Exception as exc:
        return _empty_eoi_state(user_prompt, f"EOI analysis fallback triggered: {exc}")
