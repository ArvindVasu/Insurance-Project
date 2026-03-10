from __future__ import annotations

import os
from pathlib import Path
import re

import pandas as pd
from docx import Document
from pypdf import PdfReader

from services.Common_Functions import prune_state
from services.Graph_state import GraphState
from services.llm_service import call_llm

STATE_KEYS_SET_AT_ENTRY = []
DOCUMENT_ANALYST_SYSTEM_PROMPT = (
    "You are an underwriting document analyst. Focus on insurance coverage terms, exclusions, "
    "limits, deductibles, operational risk signals, missing information, and decision-relevant "
    "underwriting implications. Be concise, structured, and practical."
)
DOCUMENT_COMPARISON_SYSTEM_PROMPT = (
    "You are an underwriting analyst comparing two insurance documents. Identify material wording "
    "changes, guideline changes, appetite changes, exclusions, deductibles, limits, authority "
    "thresholds, documentation requirements, and the underwriting impact. Answer the user's "
    "question first, then summarize the differences clearly."
)
DEFAULT_COMPARE_PROMPT = (
    "What are the key differences between the current-year and prior-year document, and what is the underwriting impact?"
)


def _extract_docx_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf_text(path: str) -> str:
    reader = PdfReader(path)
    pages = [p.extract_text() or "" for p in reader.pages]
    return "\n".join(pages)


def _load_document(path: str, is_excel: bool, is_docx: bool):
    ext = Path(path).suffix.lower()

    if is_excel or ext in {".xlsx", ".xls", ".csv"}:
        if ext == ".csv":
            df = pd.read_csv(path)
        else:
            df = pd.read_excel(path)
        preview = df.head(25)
        table_context = preview.to_markdown(index=False)
        summary = (
            f"Document type: tabular\n"
            f"Rows: {len(df)}\n"
            f"Columns: {', '.join(df.columns.astype(str).tolist())}\n"
            f"Preview:\n{table_context}"
        )
        return summary, preview

    if is_docx or ext == ".docx":
        return _extract_docx_text(path), None

    if ext == ".pdf":
        return _extract_pdf_text(path), None

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read(), None


def _truncate_context(text: str, limit: int = 12000) -> str:
    return (text or "")[:limit]


def _normalize_lines(text: str) -> list[str]:
    lines: list[str] = []
    seen: set[str] = set()
    for raw in (text or "").splitlines():
        cleaned = re.sub(r"\s+", " ", raw).strip(" -*:\t")
        if len(cleaned) < 8:
            continue
        normalized = cleaned.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        lines.append(cleaned)
    return lines


def _build_compare_fallback_summary(
    user_question: str,
    current_name: str,
    current_text: str,
    prior_name: str,
    prior_text: str,
) -> str:
    current_lines = _normalize_lines(current_text)
    prior_lines = _normalize_lines(prior_text)
    prior_norms = {line.lower() for line in prior_lines}
    current_norms = {line.lower() for line in current_lines}

    additions = [line for line in current_lines if line.lower() not in prior_norms][:6]
    removals = [line for line in prior_lines if line.lower() not in current_norms][:6]

    sections = [
        f"Answer To Question: Compared `{current_name}` against `{prior_name}` to address: {user_question or DEFAULT_COMPARE_PROMPT}",
        "Key Differences In Current-Year Document:",
    ]
    if additions:
        sections.extend([f"- {line}" for line in additions])
    else:
        sections.append("- No clear additions or tightened clauses were deterministically identified.")

    sections.append("Removed / Relaxed From Prior-Year Document:")
    if removals:
        sections.extend([f"- {line}" for line in removals])
    else:
        sections.append("- No clear removals or relaxed clauses were deterministically identified.")

    sections.extend(
        [
            "Underwriting Impact:",
            "- Review the highlighted clause changes before applying prior-year underwriting assumptions.",
            "Items To Validate:",
            "- Confirm whether any changed limits, exclusions, authority thresholds, or documentation requirements should alter the current underwriting approach.",
        ]
    )
    return "\n".join(sections)


def _build_single_document_prompt(user_prompt: str, document_text: str) -> str:
    context = _truncate_context(document_text)
    return f"""
User question:
{user_prompt}

Document content (possibly truncated):
{context}

Provide:
1) A concise answer to the user question.
2) Risk considerations and underwriting implications.
3) Missing information the underwriter should ask for next.
"""


def _build_document_comparison_prompt(
    user_prompt: str,
    current_name: str,
    current_text: str,
    prior_name: str,
    prior_text: str,
) -> str:
    compare_request = user_prompt.strip() or DEFAULT_COMPARE_PROMPT
    current_context = _truncate_context(current_text)
    prior_context = _truncate_context(prior_text)

    return f"""
You are an underwriting analyst comparing two insurance documents.

User request:
{compare_request}

Treat the first document as the current-year version and the second document as the prior-year version unless the wording clearly indicates otherwise.

Current-Year Document: {current_name}
{current_context}

Prior-Year Document: {prior_name}
{prior_context}

Provide a concise markdown response with these sections:
1. Answer To User Question
2. Key Differences In Current-Year Document
3. Removed Or Relaxed From Prior-Year Document
4. Underwriting Impact
5. Items To Validate With Broker / Underwriter

Focus on guideline changes, appetite changes, authority thresholds, coverage wording, exclusions, deductibles, limits, documentation requirements, and any underwriting actions that should change because of the differences.
"""


def document_node(state: GraphState) -> GraphState:
    file_path = state.get("uploaded_file1_path")
    if not file_path or not os.path.exists(file_path):
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "route": "document",
            "general_summary": "No document uploaded. Please attach a file to analyze.",
            "sql_result": None,
        }

    try:
        document_text_1, preview_df_1 = _load_document(
            file_path,
            bool(state.get("uploaded_file1_is_excel")),
            bool(state.get("uploaded_file1_is_docx")),
        )
    except Exception as exc:
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "route": "document",
            "general_summary": f"Failed to read document: {exc}",
            "sql_result": None,
        }

    second_file_path = state.get("uploaded_file2_path")
    if second_file_path and os.path.exists(second_file_path):
        try:
            document_text_2, preview_df_2 = _load_document(
                second_file_path,
                bool(state.get("uploaded_file2_is_excel")),
                bool(state.get("uploaded_file2_is_docx")),
            )
        except Exception as exc:
            return {
                **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
                "route": "document",
                "general_summary": f"Failed to read comparison document: {exc}",
                "comparison_summary": None,
                "sql_result": preview_df_1,
                "document_compare_preview_1": preview_df_1,
                "document_compare_preview_2": None,
                "document_compare_name_1": state.get("uploaded_file1_name") or Path(file_path).name,
                "document_compare_name_2": state.get("uploaded_file2_name") or Path(second_file_path).name,
            }

        current_name = state.get("uploaded_file1_name") or Path(file_path).name
        prior_name = state.get("uploaded_file2_name") or Path(second_file_path).name
        compare_prompt = _build_document_comparison_prompt(
            state.get("user_prompt", ""),
            current_name,
            document_text_1,
            prior_name,
            document_text_2,
        )
        comparison = call_llm(
            compare_prompt,
            system_prompt=DOCUMENT_COMPARISON_SYSTEM_PROMPT,
        ).strip()
        if not comparison or comparison.lower().startswith("openai call failed"):
            comparison = _build_compare_fallback_summary(
                state.get("user_prompt", "").strip(),
                current_name,
                document_text_1,
                prior_name,
                document_text_2,
            )

        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "route": "document",
            "general_summary": comparison,
            "comparison_summary": comparison,
            "sql_result": preview_df_1,
            "document_compare_preview_1": preview_df_1,
            "document_compare_preview_2": preview_df_2,
            "document_compare_name_1": current_name,
            "document_compare_name_2": prior_name,
        }

    insight_prompt = _build_single_document_prompt(
        state.get("user_prompt", ""),
        document_text_1,
    )
    analysis = call_llm(
        insight_prompt,
        system_prompt=DOCUMENT_ANALYST_SYSTEM_PROMPT,
    )

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "route": "document",
        "general_summary": analysis,
        "comparison_summary": None,
        "sql_result": preview_df_1,
        "document_compare_preview_1": None,
        "document_compare_preview_2": None,
        "document_compare_name_1": None,
        "document_compare_name_2": None,
    }

