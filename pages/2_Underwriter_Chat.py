from __future__ import annotations

import tempfile
import uuid
from datetime import datetime
from io import BytesIO
import os
import re
import json
import sqlite3
import hashlib
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.vanna_agent import get_user_chart_type, plot_chart, suggest_chart
from services.auth_service import require_auth
from services.Common_Functions import (
    format_date_label,
    generate_title,
    serialize_chat_history,
    _get_entry_datetime,
    safe_serialize_obj,
)
from services.Export_To_PPT import generate_ppt
from services.Follow_up_ques import generate_follow_up_questions
from services.Graph_builder import build_graph
from services.Graph_state import GraphState
from services.Output_Functions import _render_run_by_route
from services.Visualize_Workflow import visualize_workflow
from services.parse_output import parse_output
from services.ui_theme import apply_theme, render_hero, render_top_nav

require_auth()
BASE_DIR = Path(__file__).resolve().parent.parent
CHAT_DB_PATH = BASE_DIR / "chatbot.db"
CHAT_STATE_KEY_PREFIX = "underwriter_chat_state_v1"


def _db_conn() -> sqlite3.Connection:
    return sqlite3.connect(str(CHAT_DB_PATH))


def _resolve_user_id() -> str:
    # Use only login-session email for user-scoped chat storage.
    value = st.session_state.get("user_email") or st.session_state.get("email")
    if value:
        return str(value).strip().lower()

    # Fallback keeps local runs stable if login identity is missing.
    return "anonymous@local"


def _chat_state_key() -> str:
    user_id = _resolve_user_id()
    user_hash = hashlib.sha256(user_id.encode("utf-8")).hexdigest()[:16]
    return f"{CHAT_STATE_KEY_PREFIX}:{user_hash}"


def _ensure_chat_state_table() -> None:
    conn = _db_conn()
    try:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS app_kv_store (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            """
        )
        conn.commit()
    finally:
        conn.close()


def _default_current_session() -> dict:
    return {
        "id": str(uuid.uuid4()),
        "title": None,
        "created_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "messages": [],
    }


def _build_history_entry_from_session(sess: dict, archived_at: str | None = None) -> dict:
    messages = list(sess.get("messages") or [])
    last_run = messages[-1].get("assistant_run", {}) if messages else {}
    first_prompt = messages[0].get("user_prompt", "") if messages else ""
    return {
        "id": sess.get("id") or str(uuid.uuid4()),
        "title": sess.get("title") or (generate_title(first_prompt) if first_prompt else None),
        "prompt": first_prompt,
        "route": last_run.get("route"),
        "result": last_run.get("result"),
        "sql_query": last_run.get("sql_query"),
        "web_links": last_run.get("web_links"),
        "general_summary": last_run.get("general_summary"),
        "comparison_summary": last_run.get("comparison_summary"),
        "chart_info": last_run.get("chart_info"),
        "faiss_summary": last_run.get("faiss_summary"),
        "faiss_sources": last_run.get("faiss_sources"),
        "faiss_images": last_run.get("faiss_images"),
        "intranet_summary": last_run.get("intranet_summary"),
        "intranet_sources": last_run.get("intranet_sources"),
        "intranet_doc_links": last_run.get("intranet_doc_links"),
        "intranet_doc_count": last_run.get("intranet_doc_count"),
        "messages": messages,
        "created_at": sess.get("created_at"),
        "archived_at": archived_at,
    }


def _upsert_current_session_in_history(mark_archived: bool = False) -> None:
    sess = st.session_state.get("current_session") or {}
    if not (sess.get("messages") or []):
        return
    archived_at = datetime.now().strftime("%d %b %Y, %I:%M %p") if mark_archived else None
    entry = _build_history_entry_from_session(sess, archived_at=archived_at)
    target_id = entry["id"]
    for i, existing in enumerate(st.session_state.chat_history):
        if existing.get("id") == target_id:
            if not archived_at and existing.get("archived_at"):
                entry["archived_at"] = existing.get("archived_at")
            st.session_state.chat_history[i] = entry
            return
    st.session_state.chat_history.append(entry)


def _save_chat_state() -> None:
    _ensure_chat_state_table()
    state_obj = {
        "chat_history": st.session_state.get("chat_history", []),
        "active_chat_index": st.session_state.get("active_chat_index"),
        "current_session": st.session_state.get("current_session", _default_current_session()),
    }
    payload = json.dumps(safe_serialize_obj(state_obj), ensure_ascii=False)
    conn = _db_conn()
    try:
        conn.execute(
            """
            INSERT INTO app_kv_store (key, value, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(key) DO UPDATE SET
                value=excluded.value,
                updated_at=excluded.updated_at
            """,
            (_chat_state_key(), payload, datetime.now().isoformat()),
        )
        conn.commit()
    finally:
        conn.close()


def _load_chat_state() -> dict | None:
    _ensure_chat_state_table()
    conn = _db_conn()
    try:
        row = conn.execute("SELECT value FROM app_kv_store WHERE key = ?", (_chat_state_key(),)).fetchone()
    finally:
        conn.close()
    if not row or not row[0]:
        return None
    try:
        return json.loads(row[0])
    except Exception:
        return None


def _init_state_from_db_once() -> None:
    current_user = _resolve_user_id()
    loaded_for_user = st.session_state.get("_underwriter_chat_loaded_user")
    if loaded_for_user == current_user:
        return

    loaded = _load_chat_state()
    st.session_state.chat_history = (loaded or {}).get("chat_history") or []
    st.session_state.active_chat_index = (loaded or {}).get("active_chat_index")
    st.session_state.just_ran_agent = False
    st.session_state.current_session = (loaded or {}).get("current_session") or _default_current_session()

    active_idx = st.session_state.active_chat_index
    if active_idx is not None and (active_idx < 0 or active_idx >= len(st.session_state.chat_history)):
        st.session_state.active_chat_index = None

    st.session_state["_underwriter_chat_loaded_user"] = current_user


def _resume_history_session(index: int) -> None:
    if index < 0 or index >= len(st.session_state.chat_history):
        return
    entry = st.session_state.chat_history[index]
    st.session_state.current_session = {
        "id": entry.get("id") or str(uuid.uuid4()),
        "title": entry.get("title"),
        "created_at": entry.get("created_at") or datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "messages": list(entry.get("messages") or []),
    }
    st.session_state.active_chat_index = None
    st.session_state.just_ran_agent = False
    _save_chat_state()
    st.rerun()


def _to_dataframe(result):
    if isinstance(result, pd.DataFrame):
        return result
    if isinstance(result, list):
        if result and isinstance(result[0], dict):
            return pd.DataFrame(result)
        return pd.DataFrame({"Result": [str(x) for x in result]})
    if isinstance(result, dict):
        return pd.DataFrame([result])
    if result is None:
        return None
    return pd.DataFrame({"Result": [str(result)]})


def _add_dataframe_table(doc: Document, df: pd.DataFrame, title: str = "Tabular Result") -> None:
    if df is None or df.empty:
        return
    doc.add_paragraph(title).runs[0].bold = True
    # Keep docs readable for larger outputs.
    df = df.head(30)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row.tolist()):
            cells[i].text = "" if pd.isna(val) else str(val)


def _extract_title_url(link_md: str) -> tuple[str, str]:
    match = re.match(r"\[(.*?)\]\((.*?)\)", str(link_md or "").strip())
    if match:
        return match.group(1), match.group(2)
    return "Link", str(link_md or "")


def _add_web_links(doc: Document, web_links) -> None:
    if not web_links:
        return
    doc.add_paragraph("Web Links").runs[0].bold = True
    for i, item in enumerate(web_links, start=1):
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            link_md, summary = item[0], item[1]
        else:
            link_md, summary = item, ""
        title, url = _extract_title_url(link_md)
        doc.add_paragraph(f"{i}. {title}")
        doc.add_paragraph(f"URL: {url}")
        if summary:
            doc.add_paragraph(f"Summary: {summary}")


def _add_images(doc: Document, images_meta) -> None:
    if not images_meta:
        return
    valid = [m for m in images_meta if m.get("extracted_image_path") and os.path.exists(m.get("extracted_image_path"))]
    if not valid:
        return
    doc.add_paragraph("Images").runs[0].bold = True
    for meta in valid[:8]:
        caption = meta.get("caption") or meta.get("original_doc") or "Image"
        doc.add_paragraph(str(caption))
        try:
            doc.add_picture(meta["extracted_image_path"], width=Inches(5.8))
        except Exception:
            doc.add_paragraph(f"[Could not embed image: {meta['extracted_image_path']}]")


def generate_doc(entry: dict) -> BytesIO:
    doc = Document()
    logo = doc.add_paragraph("ASTRA")
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo.runs[0]
    logo_run.bold = True
    logo_run.font.size = doc.styles["Title"].font.size
    subtitle = doc.add_paragraph("AI Underwriting Assistant")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Agentic AI Report", level=1)
    title = entry.get("title") or entry.get("prompt") or "Session"
    doc.add_paragraph(f"Session: {title}")
    if entry.get("created_at"):
        doc.add_paragraph(f"Created: {entry['created_at']}")

    for idx, turn in enumerate(entry.get("messages", []), start=1):
        doc.add_heading(f"Turn {idx}", level=2)
        user_prompt = turn.get("user_prompt") or ""
        doc.add_paragraph(f"You: {user_prompt}")
        assistant_run = turn.get("assistant_run") or {}
        if assistant_run:
            route = assistant_run.get("route")
            if route:
                doc.add_paragraph(f"Route: {route}")
            if assistant_run.get("sql_query"):
                doc.add_paragraph("SQL Query:")
                doc.add_paragraph(str(assistant_run.get("sql_query")))
            if assistant_run.get("comparison_summary"):
                doc.add_paragraph("Comparison Summary:")
                doc.add_paragraph(str(assistant_run.get("comparison_summary")))
            if assistant_run.get("general_summary"):
                doc.add_paragraph("Summary:")
                doc.add_paragraph(str(assistant_run.get("general_summary")))
            result_df = _to_dataframe(assistant_run.get("result"))
            _add_dataframe_table(doc, result_df, title="Tabular Result")

            # search/comp links
            web_links = assistant_run.get("web_links")
            if route == "search" and not web_links:
                web_links = assistant_run.get("result")
            _add_web_links(doc, web_links)

            # intranet links (if any)
            intranet_links = assistant_run.get("intranet_doc_links") or []
            if intranet_links:
                doc.add_paragraph("Intranet Document Links").runs[0].bold = True
                for i, link in enumerate(intranet_links, start=1):
                    doc.add_paragraph(f"{i}. {link}")

            # faiss images (if any)
            _add_images(doc, assistant_run.get("faiss_images") or [])
        if turn.get("timestamp"):
            doc.add_paragraph(f"Time: {turn.get('timestamp')}")

    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff

apply_theme("Underwriter Chat", icon=":speech_balloon:")
render_top_nav(show_search=False)
render_hero(
    "Underwriter Chat",
    "Original ASTRA routed workflow: router -> specialized agent nodes (SQL/Search/Comp/FAISS/Intranet/Document).",
)
st.markdown(
    """
    <style>
    .uw-form-card {
      background: linear-gradient(180deg, #ffffff 0%, #f5f9ff 100%);
      border: 2px solid #355f8f;
      border-radius: 14px;
      padding: 12px 14px 2px 14px;
      margin: 10px 0 12px 0;
      box-shadow: 0 10px 22px rgba(18, 42, 76, 0.12);
    }

    .uw-form-card label, .uw-form-card [data-testid="stWidgetLabel"] {
      color: #12335f !important;
      font-weight: 700 !important;
      font-size: 16px !important;
      letter-spacing: 0.2px;
    }

    [data-testid="stSpinner"] {
      background: linear-gradient(90deg, #eaf5ff 0%, #eefaf8 100%);
      border: 1px solid #b9d9e8;
      border-radius: 999px;
      padding: 8px 14px;
      width: fit-content;
      margin-top: 8px;
      box-shadow: 0 4px 10px rgba(16, 64, 90, 0.08);
    }

    [data-testid="stSpinner"] p {
      color: #124a5e;
      font-weight: 700;
      margin: 0;
    }

    /* Make query box visually prominent */
    [data-testid="stTextInput"] input {
      background: #ffffff !important;
      border: 2px solid #7aa5d8 !important;
      border-radius: 12px !important;
      color: #0f2d52 !important;
      box-shadow: 0 4px 12px rgba(30, 80, 140, 0.08);
      font-weight: 600;
    }
    [data-testid="stTextInput"] input:focus {
      border-color: #0f766e !important;
      box-shadow: 0 0 0 3px rgba(15, 118, 110, 0.18), 0 6px 14px rgba(15, 118, 110, 0.12) !important;
    }
    [data-testid="stTextInput"] input::placeholder {
      color: #5e7ea8;
      opacity: 1;
    }

    /* Run Agent button */
    .uw-form-card [data-testid="stFormSubmitButton"] button,
    .uw-form-card .stButton button {
      background: linear-gradient(135deg, #0b4f91 0%, #0f766e 100%) !important;
      color: #ffffff !important;
      border: 1px solid #0b4f91 !important;
      border-radius: 10px !important;
      font-weight: 700 !important;
      font-size: 18px !important;
      min-height: 44px;
      box-shadow: 0 8px 18px rgba(11, 79, 145, 0.24);
      transition: transform 0.15s ease, box-shadow 0.2s ease;
    }
    .uw-form-card [data-testid="stFormSubmitButton"] button:hover,
    .uw-form-card .stButton button:hover {
      transform: translateY(-1px);
      box-shadow: 0 10px 22px rgba(11, 79, 145, 0.3);
    }
    </style>
    """,
    unsafe_allow_html=True,
)

_init_state_from_db_once()

with st.sidebar:
    st.subheader("Sessions")
    if st.button("Start New Session", use_container_width=True):
        _upsert_current_session_in_history(mark_archived=True)

        st.session_state.current_session = {
            "id": str(uuid.uuid4()),
            "title": None,
            "created_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
            "messages": [],
        }
        st.session_state.active_chat_index = None
        _save_chat_state()
        st.rerun()

    if st.button("Clear All History", use_container_width=True):
        st.session_state.chat_history = []
        st.session_state.active_chat_index = None
        st.session_state.current_session = _default_current_session()
        _save_chat_state()

    history_json = serialize_chat_history(st.session_state.chat_history)
    st.download_button(
        "Export History",
        history_json,
        file_name="chat_history.json",
        use_container_width=True,
    )

    grouped = {}
    for chat in st.session_state.chat_history:
        chat_date = _get_entry_datetime(chat).date()
        grouped.setdefault(chat_date, []).append(chat)

    for group_date in sorted(grouped.keys(), reverse=True):
        with st.expander(format_date_label(group_date), expanded=False):
            for idx, chat in enumerate(grouped[group_date]):
                title = chat.get("title") or chat.get("prompt", "Session")[:40]
                if st.button(title, key=f"hist_{group_date}_{idx}", use_container_width=True):
                    resume_idx = st.session_state.chat_history.index(chat)
                    _resume_history_session(resume_idx)

if st.session_state.active_chat_index is None:
    st.markdown('<div class="uw-form-card">', unsafe_allow_html=True)
    with st.form("underwriter_chat_form"):
        c1, c2 = st.columns([4, 1])
        with c1:
            user_prompt = st.text_input("Enter your query", placeholder="Ask about premiums, loss ratio, IBNR, policy docs, or benchmarks")
        with c2:
            uploaded_file1 = st.file_uploader(
                "Attach file",
                type=["docx", "xlsx", "xls", "csv", "pdf", "txt"],
                label_visibility="collapsed",
            )

        submitted = st.form_submit_button("Run Agent", use_container_width=True, type="primary")
    st.markdown("</div>", unsafe_allow_html=True)
    st.caption("Tip: include context like line of business, year range, and metric for sharper answers.")

    if submitted:
        uploaded_file1_path = None
        uploaded_file1_is_excel = False
        uploaded_file1_is_docx = False

        if uploaded_file1 is not None:
            ext1 = uploaded_file1.name.split(".")[-1].lower().strip()
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext1}") as tmp1:
                tmp1.write(uploaded_file1.read())
                uploaded_file1_path = tmp1.name
            uploaded_file1_is_excel = ext1 in ["xls", "xlsx", "csv"]
            uploaded_file1_is_docx = ext1 in ["docx"]

        state: GraphState = {
            "user_prompt": user_prompt,
            "uploaded_file1": uploaded_file1,
            "uploaded_file1_path": uploaded_file1_path,
            "uploaded_file1_is_excel": uploaded_file1_is_excel,
            "uploaded_file1_is_docx": uploaded_file1_is_docx,
            "vanna_prompt": None,
            "fuzzy_prompt": None,
            "route": None,
            "sql_result": None,
            "sql_query": None,
            "web_links": None,
            "comparison_summary": None,
            "general_summary": None,
            "chart_info": None,
            "faiss_summary": None,
            "faiss_sources": None,
            "faiss_images": None,
            "intranet_summary": None,
            "intranet_sources": None,
            "intranet_doc_links": None,
            "intranet_doc_count": None,
        }

        progress_text = st.empty()
        progress_bar = st.progress(0, text="Starting...")

        with st.status("Running underwriter agent...", expanded=True) as run_status:
            progress_text.info("1/3 Routing")
            progress_bar.progress(33, text="1/3 Routing request to the right workflow...")
            run_status.write("1/3 Routing request to the right workflow...")
            agent_graph = build_graph()

            progress_text.info("2/3 Fetching")
            progress_bar.progress(66, text="2/3 Fetching data, search context, and agent outputs...")
            run_status.write("2/3 Fetching data, search context, and agent outputs...")
            output = agent_graph.invoke(state)

            progress_text.info("3/3 Summarizing")
            progress_bar.progress(90, text="3/3 Summarizing and formatting final response...")
            run_status.write("3/3 Summarizing and formatting final response...")
            run_record = parse_output(user_prompt, output)
            progress_bar.progress(100, text="Completed")
            progress_text.success("Completed")
            run_status.update(label="Run completed", state="complete")

        st.session_state.current_session["messages"].append(
            {
                "role": "turn",
                "user_prompt": user_prompt,
                "assistant_run": run_record,
                "timestamp": run_record["timestamp"],
            }
        )
        _upsert_current_session_in_history(mark_archived=False)
        _save_chat_state()
        st.session_state.followups = generate_follow_up_questions(user_prompt)

        left, right = st.columns([3.6, 1.4])
        with right:
            st.markdown("### Workflow")
            visualize_workflow(agent_graph, active_route=output.get("route"))

        with left:
            for turn in reversed(st.session_state.current_session["messages"]):
                assistant_run = turn.get("assistant_run")
                if assistant_run:
                    _render_run_by_route(assistant_run)

                    prompt = (turn.get("user_prompt") or "").strip().lower()
                    plotting_keywords = ["plot", "draw", "visualize", "chart", "bar graph", "line graph", "pie chart", "graph"]
                    if any(word in prompt for word in plotting_keywords):
                        sql_df = None
                        res = assistant_run.get("result")
                        if isinstance(res, list):
                            try:
                                sql_df = pd.DataFrame(res)
                            except Exception:
                                sql_df = None
                        elif isinstance(res, pd.DataFrame):
                            sql_df = res

                        if sql_df is not None and not sql_df.empty:
                            user_chart_type = get_user_chart_type(prompt)
                            chart_info = suggest_chart(sql_df)
                            if chart_info and user_chart_type:
                                chart_info["type"] = user_chart_type
                            if chart_info:
                                plot_chart(sql_df, chart_info)

                if turn.get("user_prompt"):
                    st.markdown(f"**You:** {turn['user_prompt']}")
                st.caption(turn.get("timestamp", ""))
                st.markdown("---")

            if st.session_state.get("followups"):
                st.markdown("### Follow-up prompts")
                for q in st.session_state["followups"]:
                    st.markdown(f"- {q}")

            try:
                sess = st.session_state.get("current_session", {})
                entry_for_export = {
                    "id": sess.get("id"),
                    "title": sess.get("title"),
                    "prompt": sess.get("messages")[0]["user_prompt"] if sess.get("messages") else "",
                    "messages": sess.get("messages", []),
                    "created_at": sess.get("created_at"),
                }
                if entry_for_export["messages"]:
                    ppt_buffer = generate_ppt(entry_for_export)
                    doc_buffer = generate_doc(entry_for_export)
                    c_ppt, c_doc = st.columns(2)
                    with c_ppt:
                        st.download_button(
                            label="Export to PPT",
                            data=ppt_buffer,
                            file_name="underwriter_chat_session.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                            key=f"download_ppt_live_{entry_for_export.get('id')}",
                            use_container_width=True,
                        )
                    with c_doc:
                        st.download_button(
                            label="Export to DOC",
                            data=doc_buffer,
                            file_name="underwriter_chat_session.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_doc_live_{entry_for_export.get('id')}",
                            use_container_width=True,
                        )
            except Exception as e:
                st.warning(f"PPT export not available: {e}")
    elif st.session_state.current_session.get("messages"):
        left, right = st.columns([3.6, 1.4])
        with right:
            st.markdown("### Workflow")
            try:
                last_route = (
                    st.session_state.current_session["messages"][-1]
                    .get("assistant_run", {})
                    .get("route")
                )
                visualize_workflow(build_graph(), active_route=last_route)
            except Exception:
                st.info("Workflow view unavailable for this session.")

        with left:
            for turn in reversed(st.session_state.current_session["messages"]):
                assistant_run = turn.get("assistant_run")
                if assistant_run:
                    _render_run_by_route(assistant_run)
                if turn.get("user_prompt"):
                    st.markdown(f"**You:** {turn['user_prompt']}")
                st.caption(turn.get("timestamp", ""))
                st.markdown("---")

            try:
                sess = st.session_state.get("current_session", {})
                entry_for_export = {
                    "id": sess.get("id"),
                    "title": sess.get("title"),
                    "prompt": sess.get("messages")[0]["user_prompt"] if sess.get("messages") else "",
                    "messages": sess.get("messages", []),
                    "created_at": sess.get("created_at"),
                }
                if entry_for_export["messages"]:
                    ppt_buffer = generate_ppt(entry_for_export)
                    doc_buffer = generate_doc(entry_for_export)
                    c_ppt, c_doc = st.columns(2)
                    with c_ppt:
                        st.download_button(
                            label="Export to PPT",
                            data=ppt_buffer,
                            file_name="underwriter_chat_session.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                            key=f"download_ppt_live_{entry_for_export.get('id')}",
                            use_container_width=True,
                        )
                    with c_doc:
                        st.download_button(
                            label="Export to DOC",
                            data=doc_buffer,
                            file_name="underwriter_chat_session.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_doc_live_{entry_for_export.get('id')}",
                            use_container_width=True,
                        )
            except Exception as e:
                st.warning(f"PPT export not available: {e}")

else:
    entry = st.session_state.chat_history[st.session_state.active_chat_index]
    st.markdown(f"### Session: {entry.get('title') or entry.get('prompt')}")
    if entry.get("created_at"):
        st.caption(f"Created: {entry['created_at']}")
    if entry.get("archived_at"):
        st.caption(f"Archived: {entry['archived_at']}")
    if st.button("Resume This Session", use_container_width=False):
        _resume_history_session(st.session_state.active_chat_index)

    for idx, turn in enumerate(entry.get("messages", []), start=1):
        st.markdown(f"**{idx}. You:** {turn.get('user_prompt', '')}")
        assistant_run = turn.get("assistant_run")
        if assistant_run:
            _render_run_by_route(assistant_run)
        st.caption(turn.get("timestamp", ""))
        st.markdown("---")

    try:
        if entry.get("messages"):
            ppt_buffer = generate_ppt(entry)
            doc_buffer = generate_doc(entry)
            c_ppt, c_doc = st.columns(2)
            with c_ppt:
                st.download_button(
                    "Export to PPT",
                    ppt_buffer,
                    file_name="underwriter_chat_history.pptx",
                    mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                    use_container_width=True,
                    key=f"download_ppt_history_{entry.get('id', 'session')}",
                )
            with c_doc:
                st.download_button(
                    "Export to DOC",
                    doc_buffer,
                    file_name="underwriter_chat_history.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"download_doc_history_{entry.get('id', 'session')}",
                )
    except Exception as e:
        st.warning(f"Unable to export PPT for this session: {e}")
